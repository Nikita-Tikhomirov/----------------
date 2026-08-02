from __future__ import annotations

import hashlib
import html
import json
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Sequence

from PIL import Image as PILImage
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader
try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        Image,
        LongTable,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    REPORTLAB_AVAILABLE = True
except ModuleNotFoundError:
    REPORTLAB_AVAILABLE = False


ROOT = Path(__file__).resolve().parents[1]
REPORT_DATE = "02.08.2026"

OUTPUT_DOCX = ROOT / "output/documents/AP-Real-form-corrections-final-2026-08-02.docx"
OUTPUT_PDF = ROOT / "output/pdf/AP-Real-form-corrections-final-2026-08-02.pdf"
OUTPUT_EMAIL = ROOT / "output/AP-Real-form-corrections-letter-2026-08-02.txt"
OUTPUT_AUDIT = ROOT / "output/ap-real-conflict-closure-audit-2026-08-02.json"

VISUAL_QA = ROOT / "output/ap-real-conflict-closure-qa-2026-08-02/summary.json"
RECIPIENT_MATRIX = ROOT / "output/ap-real-live-recipient-matrix-2026-08-02.json"
DELIVERY_QA = ROOT / "output/ap-real-form-delivery-2026-08-02.json"
ROUTE_RETEST = ROOT / "output/ap-real-client-route-restore-delivery-2026-08-02.json"
CF7_VERIFY = ROOT / "output/ap-real-cf7-recipient-normalization-verify-fresh-2026-08-02.json"
MX_AUDIT = ROOT / "output/ap-real-recipient-mx-audit-2026-08-02.json"
MAILBOX_PROOF = (
    ROOT
    / "output/ap-real-conflict-closure-evidence-2026-08-02/mailru-medlic-final-route.json"
)
MAILBOX_SCREENSHOT = (
    ROOT
    / "output/ap-real-conflict-closure-evidence-2026-08-02/mailru-medlic-final-route.png"
)
VISUAL_SHEETS = ROOT / "output/ap-real-conflict-closure-visual-review-2026-08-02/sheets"

INCLUDED_SITES = (
    "apreal.ru",
    "docp.ru",
    "mca24.ru",
    "fsa-lab.ru",
    "elecktro.ru",
    "med-license.ru",
    "mhsl.ru",
    "fste.ru",
    "otxodi.ru",
    "lfsb.ru",
    "apreal.spb.ru",
    "minkult78.ru",
    "medlic.spb.ru",
    "medtex78.ru",
    "mchs-spb.ru",
    "mchs78.ru",
    "license39.ru",
    "medtex39.ru",
    "39mchs.ru",
    "apreal36.ru",
    "apreal-nn.ru",
    "apreal-volgograd.ru",
    "apreal72.ru",
    "nousro.ru",
    "dpomuc.ru",
    "nousro-spb.ru",
    "ed-kgd.ru",
    "muc-vrn.ru",
    "nousro-nn.ru",
    "shopap.ru",
)

EXCLUDED_SITES = (
    "rectavr.ru",
    "fstek.spb.ru",
    "lic-k.ru",
    "apreal-samara.ru",
    "ed-krd.ru",
)

CLIENT_SECTION_TITLES = (
    "Что требовалось исправить",
    "Что сделано",
    "Результаты повторной проверки",
    "Матрица сайтов и получателей",
    "Визуальные подтверждения",
)

REQUIREMENTS = (
    (
        "Две формы на каждом сайте",
        "Доступны формы «ЗАКАЗАТЬ ЗВОНОК» и «ЗАДАТЬ ВОПРОС».",
    ),
    (
        "Форма заказа звонка",
        "Имя необязательно, телефон обязателен, капча есть; email и тип клиента убраны.",
    ),
    (
        "Форма вопроса",
        "Имя и вопрос необязательны, телефон обязателен, капча есть; email убран.",
    ),
    (
        "Согласие и политика",
        "Согласованный текст показан без отдельной галочки; ссылка ведёт на политику конфиденциальности.",
    ),
    (
        "Работа окон",
        "Крестик виден, окна закрываются и не перекрываются другими элементами.",
    ),
    (
        "Корректная отправка",
        "Нет пустой страницы, 404, вечной загрузки и ложного успеха; подтверждение появляется после принятия заявки.",
    ),
    (
        "Почтовые получатели",
        "Каждый сайт отправляет заявку на адрес, указанный для него в исходной таблице клиента.",
    ),
    (
        "Повторная проверка",
        "Каждый сайт и обе формы проверены после исправлений на компьютере и телефоне.",
    ),
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E0E7"
DARK_GRAY = "4B5563"
GREEN = "1F7A45"
PALE_GREEN = "E8F5EC"
BLACK = "111111"

if REPORTLAB_AVAILABLE:
    PDF_BLUE = colors.HexColor(f"#{BLUE}")
    PDF_DARK_BLUE = colors.HexColor(f"#{DARK_BLUE}")
    PDF_PALE_BLUE = colors.HexColor(f"#{PALE_BLUE}")
    PDF_LIGHT_GRAY = colors.HexColor(f"#{LIGHT_GRAY}")
    PDF_MID_GRAY = colors.HexColor(f"#{MID_GRAY}")
    PDF_DARK_GRAY = colors.HexColor(f"#{DARK_GRAY}")
    PDF_GREEN = colors.HexColor(f"#{GREEN}")
    PDF_PALE_GREEN = colors.HexColor(f"#{PALE_GREEN}")
    PDF_BLACK = colors.HexColor(f"#{BLACK}")
else:
    PDF_BLUE = PDF_DARK_BLUE = PDF_PALE_BLUE = None
    PDF_LIGHT_GRAY = PDF_MID_GRAY = PDF_DARK_GRAY = None
    PDF_GREEN = PDF_PALE_GREEN = PDF_BLACK = None


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8-sig"))


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def summarize_visual_qa(
    records: Iterable[dict[str, Any]], included_sites: set[str]
) -> dict[str, Any]:
    included = [record for record in records if record.get("domain") in included_sites]
    failures = [
        {
            "domain": record.get("domain"),
            "viewport": record.get("viewport"),
            "failures": record.get("failures", []),
        }
        for record in included
        if record.get("status") != 200 or record.get("failures")
    ]
    return {
        "sites": len({record.get("domain") for record in included}),
        "views": len(included),
        "failed_views": len(failures),
        "failures": failures,
    }


def build_client_email_text() -> str:
    return """Альберт, добрый день.

Предыдущий отчёт прошу не учитывать. В нём я преждевременно свёл результаты точечных проверок и не развернул повторяющиеся замечания на весь список однотипных сайтов. Из-за этого часть работ была ошибочно отмечена как завершённая. Это моя ошибка.

Я заново восстановил требования из присланных файлов и перепроверил весь спорный пакет по полной матрице: 30 сайтов, две формы на каждом, компьютерная и мобильная версии, обработчики заявок и почтовые получатели. Недостающие настройки получателей возвращены на штатные адреса, после чего изменённые маршруты проверены повторной отправкой.

Во вложении короткий итоговый отчёт только по этому пакету. В нём перечислены требования, показана матрица всех 30 сайтов и приложены визуальные подтверждения. Предыдущий отчёт прошу не использовать.

С уважением,
Никита
"""


def visual_sheet_paths() -> list[Path]:
    paths = sorted(VISUAL_SHEETS.glob("sheet-*.png"))
    if len(paths) != 10:
        raise ValueError(f"Expected 10 visual evidence sheets, got {len(paths)}")
    return paths


def validate_sources() -> dict[str, Any]:
    required_paths = (
        VISUAL_QA,
        RECIPIENT_MATRIX,
        DELIVERY_QA,
        ROUTE_RETEST,
        CF7_VERIFY,
        MX_AUDIT,
        MAILBOX_PROOF,
        MAILBOX_SCREENSHOT,
    )
    missing = [str(path) for path in required_paths if not path.exists()]
    if missing:
        raise FileNotFoundError(f"Missing report evidence: {missing}")

    visual = summarize_visual_qa(load_json(VISUAL_QA), set(INCLUDED_SITES))
    if visual != {"sites": 30, "views": 60, "failed_views": 0, "failures": []}:
        raise ValueError(f"Visual QA is not complete: {visual}")

    matrix = load_json(RECIPIENT_MATRIX)
    matrix_sites = matrix.get("sites", [])
    matrix_domains = {item.get("domain") for item in matrix_sites}
    if matrix_domains != set(INCLUDED_SITES):
        raise ValueError("Recipient matrix does not match the 30-site scope")
    if matrix.get("summary", {}).get("passed") != 30 or not all(
        item.get("passed") for item in matrix_sites
    ):
        raise ValueError("Recipient matrix contains a failed site")

    delivery = load_json(DELIVERY_QA).get("submissions", [])
    accepted_delivery = [item for item in delivery if item.get("accepted")]
    delivery_pairs = {(item.get("domain"), item.get("kind")) for item in accepted_delivery}
    expected_pairs = {
        (domain, kind)
        for domain in INCLUDED_SITES
        for kind in ("callback", "question")
    }
    if delivery_pairs != expected_pairs:
        raise ValueError("Functional form delivery evidence is incomplete")

    route_retest = load_json(ROUTE_RETEST).get("submissions", [])
    if len(route_retest) != 18 or not all(item.get("accepted") for item in route_retest):
        raise ValueError("Recipient route retest is incomplete")

    cf7 = load_json(CF7_VERIFY)
    if not cf7.get("ok") or cf7.get("verified") != 16:
        raise ValueError("Legacy Contact Form 7 verification is incomplete")

    mx = load_json(MX_AUDIT)
    mx_summary = mx.get("summary", {})
    if mx_summary.get("passed") != 30:
        raise ValueError("MX audit is incomplete")

    mailbox = load_json(MAILBOX_PROOF)
    if mailbox.get("mailbox") not in mailbox.get("to", []):
        raise ValueError("Mailbox proof does not show delivery to the target mailbox")

    sheets = visual_sheet_paths()
    return {
        "visual": visual,
        "recipient_matrix": matrix,
        "functional_submissions": len(accepted_delivery),
        "route_retest_submissions": len(route_retest),
        "cf7_forms": cf7.get("verified"),
        "mx_sites": mx_summary.get("passed"),
        "mailbox": mailbox,
        "sheets": sheets,
    }


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = "Arial"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, value: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = MID_GRAY, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char, instr_text, fld_end))


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    run_properties.extend((fonts, color, underline))
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(run_properties)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Heading 1", 16, DARK_BLUE),
        ("Heading 2", 12, BLUE),
    ):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_run_font(
        header.add_run("ГК «АП-Риал» | Итоговая проверка форм"),
        size=8,
        bold=True,
        color=DARK_GRAY,
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("Страница "), size=8, color=DARK_GRAY)
    add_field(footer, "PAGE")


def add_docx_callout(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_borders(table, "A7D5B7", 7)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GREEN)
    set_cell_margins(cell, 135)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.add_run(text), size=10, bold=True, color=GREEN)


def add_docx_metrics(document: Document) -> None:
    metrics = (
        ("30/30", "сайтов в матрице"),
        ("60/60", "обработчиков приняли заявки"),
        ("60/60", "просмотров: компьютер/телефон"),
        ("30/30", "маршрутов получателей"),
    )
    table = document.add_table(rows=1, cols=4)
    set_table_borders(table, "B8D4E8", 7)
    for cell, (value, label) in zip(table.rows[0].cells, metrics):
        set_cell_width(cell, 1.75)
        set_cell_shading(cell, PALE_BLUE)
        set_cell_margins(cell, 110)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(value), size=15, bold=True, color=DARK_BLUE)
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(label), size=7.5, color=DARK_GRAY)


def add_docx_requirements(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    set_table_borders(table)
    widths = (1.65, 4.85, 0.75)
    header = table.rows[0]
    repeat_table_header(header)
    for cell, text, width in zip(header.cells, ("Пункт", "Критерий", "Статус"), widths):
        set_cell_width(cell, width)
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        set_run_font(cell.paragraphs[0].add_run(text), size=8.5, bold=True, color=DARK_BLUE)
    for title, criterion in REQUIREMENTS:
        row = table.add_row()
        keep_row_together(row)
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_run_font(row.cells[0].paragraphs[0].add_run(title), size=8.2, bold=True)
        set_run_font(row.cells[1].paragraphs[0].add_run(criterion), size=8.2)
        set_cell_shading(row.cells[2], PALE_GREEN)
        set_run_font(row.cells[2].paragraphs[0].add_run("Готово"), size=8.2, bold=True, color=GREEN)


def add_docx_matrix(document: Document, matrix: dict[str, Any]) -> None:
    rows = sorted(matrix["sites"], key=lambda item: INCLUDED_SITES.index(item["domain"]))
    table = document.add_table(rows=1, cols=4)
    set_table_borders(table)
    widths = (0.45, 2.0, 2.9, 1.9)
    header = table.rows[0]
    repeat_table_header(header)
    for cell, text, width in zip(
        header.cells,
        ("№", "Сайт", "Получатель", "Результат"),
        widths,
    ):
        set_cell_width(cell, width)
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell, 80)
        set_run_font(cell.paragraphs[0].add_run(text), size=8, bold=True, color=DARK_BLUE)
    for index, item in enumerate(rows, start=1):
        row = table.add_row()
        keep_row_together(row)
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell, 75)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_run_font(row.cells[0].paragraphs[0].add_run(str(index)), size=7.7)
        add_hyperlink(
            row.cells[1].paragraphs[0],
            item["domain"],
            f"https://{item['domain']}/",
        )
        set_run_font(
            row.cells[2].paragraphs[0].add_run(item["expected_recipient"]),
            size=7.7,
        )
        set_cell_shading(row.cells[3], PALE_GREEN)
        set_run_font(
            row.cells[3].paragraphs[0].add_run("Обе формы проверены"),
            size=7.4,
            bold=True,
            color=GREEN,
        )


def docx_add_picture_fit(paragraph, path: Path, max_width: float, max_height: float) -> None:
    with PILImage.open(path) as image:
        width_px, height_px = image.size
    scale = min(max_width / width_px, max_height / height_px)
    paragraph.add_run().add_picture(
        str(path),
        width=Inches(width_px * scale),
        height=Inches(height_px * scale),
    )


def build_docx(evidence: dict[str, Any]) -> None:
    document = Document()
    configure_docx(document)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    set_run_font(paragraph.add_run("ИТОГОВЫЙ ОТЧЁТ"), size=10, bold=True, color=BLUE)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    set_run_font(
        paragraph.add_run("Исправление и повторная проверка форм"),
        size=23,
        bold=True,
        color=DARK_BLUE,
    )
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(14)
    set_run_font(
        paragraph.add_run("Спорный пакет работ по 30 сайтам ГК «АП-Риал»"),
        size=12,
        color=DARK_GRAY,
    )
    add_docx_callout(
        document,
        "Повторная проверка завершена по полной матрице. На 30 включённых сайтах несоответствий не осталось.",
    )
    document.add_paragraph()
    add_docx_metrics(document)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    set_run_font(paragraph.add_run(f"Дата проверки: {REPORT_DATE}"), size=9, color=DARK_GRAY)
    set_run_font(
        document.add_paragraph().add_run(
            "Объём отчёта: только последние исправления форм, почтовых маршрутов и повторная приёмка."
        ),
        size=9,
        color=DARK_GRAY,
    )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(CLIENT_SECTION_TITLES[0], level=1)
    document.add_paragraph(
        "Замечания были развернуты на весь список однотипных сайтов, а не только на отдельные примеры. "
        "Ниже приведены критерии, по которым выполнена повторная приёмка."
    )
    add_docx_requirements(document)
    paragraph = document.add_paragraph()
    set_run_font(
        paragraph.add_run(
            "Не входили в этот пакет по исходному указанию клиента: "
            + ", ".join(EXCLUDED_SITES)
            + "."
        ),
        size=8.5,
        color=DARK_GRAY,
    )

    document.add_heading(CLIENT_SECTION_TITLES[1], level=1)
    completed = (
        "Единые требования к двум формам применены ко всем 30 сайтам из матрицы.",
        "Исправлена логика отправки: подтверждение показывается только после ответа обработчика.",
        "Почтовые получатели сверены с исходной таблицей; изменённые маршруты возвращены на штатные адреса.",
        "После изменения маршрутов обе формы на затронутых сайтах отправлены повторно.",
        "Каждый сайт заново открыт и просмотрен на компьютере и телефоне; обе формы проверены отдельно.",
    )
    for item in completed:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(CLIENT_SECTION_TITLES[2], level=1)
    add_docx_metrics(document)
    results = (
        "60 из 60 обработчиков приняли валидную контрольную заявку и вернули согласованное подтверждение.",
        "18 из 18 форм на сайтах с изменённым маршрутом повторно приняли заявку после исправления адресов.",
        "30 из 30 актуальных конфигураций отправки совпадают с адресами из клиентской матрицы.",
        "16 дополнительных форм Contact Form 7 на семи сайтах отдельно проверены: получатели указаны корректно.",
        "Почтовые домены всех 30 получателей имеют рабочие MX-записи.",
        "Для сквозного контроля в доступном целевом ящике info@medlic.spb.ru подтверждены оба тестовых письма.",
        "60 из 60 контрольных просмотров страниц (30 сайтов × компьютер/телефон) завершены без замечаний; проверено 120 окон форм.",
    )
    for item in results:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)
    add_docx_callout(
        document,
        "Итог: по включённому списку не осталось незакрытых пунктов. Дополнительные данные от клиента для этого пакета не требуются.",
    )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(CLIENT_SECTION_TITLES[3], level=1)
    document.add_paragraph(
        "В таблице указан фактический получатель, прочитанный из рабочей конфигурации каждого сайта после исправлений."
    )
    add_docx_matrix(document, evidence["recipient_matrix"])

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(CLIENT_SECTION_TITLES[4], level=1)
    document.add_paragraph(
        "Каждый лист ниже содержит три сайта. Для каждого сайта показаны обе формы в компьютерной и мобильной версиях. "
        "Снимки сделаны после окончательной настройки почтовых маршрутов."
    )
    for index, sheet in enumerate(evidence["sheets"], start=1):
        if index > 1:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        set_run_font(
            paragraph.add_run(f"Лист {index} из 10"),
            size=9,
            bold=True,
            color=DARK_BLUE,
        )
        picture = document.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.paragraph_format.space_after = Pt(0)
        docx_add_picture_fit(picture, sheet, 7.0, 8.95)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)


def register_pdf_fonts() -> None:
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("ReportLab is required to build the PDF")
    regular = Path("C:/Windows/Fonts/arial.ttf")
    bold = Path("C:/Windows/Fonts/arialbd.ttf")
    if not regular.exists() or not bold.exists():
        raise FileNotFoundError("Arial fonts are required for the Cyrillic PDF")
    if "APClosureArial" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("APClosureArial", str(regular)))
        pdfmetrics.registerFont(TTFont("APClosureArial-Bold", str(bold)))
        pdfmetrics.registerFontFamily(
            "APClosureArial",
            normal="APClosureArial",
            bold="APClosureArial-Bold",
            italic="APClosureArial",
            boldItalic="APClosureArial-Bold",
        )


def pdf_styles() -> dict[str, ParagraphStyle]:
    register_pdf_fonts()
    sample = getSampleStyleSheet()
    return {
        "body": ParagraphStyle(
            "ClosureBody",
            parent=sample["BodyText"],
            fontName="APClosureArial",
            fontSize=9.2,
            leading=12.4,
            textColor=PDF_BLACK,
            spaceAfter=5,
        ),
        "small": ParagraphStyle(
            "ClosureSmall",
            parent=sample["BodyText"],
            fontName="APClosureArial",
            fontSize=7.2,
            leading=9.2,
            textColor=PDF_DARK_GRAY,
        ),
        "table": ParagraphStyle(
            "ClosureTable",
            parent=sample["BodyText"],
            fontName="APClosureArial",
            fontSize=7.2,
            leading=9.0,
            textColor=PDF_BLACK,
        ),
        "table_header": ParagraphStyle(
            "ClosureTableHeader",
            parent=sample["BodyText"],
            fontName="APClosureArial-Bold",
            fontSize=7.4,
            leading=9.2,
            textColor=PDF_DARK_BLUE,
        ),
        "kicker": ParagraphStyle(
            "ClosureKicker",
            parent=sample["BodyText"],
            fontName="APClosureArial-Bold",
            fontSize=9,
            leading=11,
            textColor=PDF_BLUE,
            spaceAfter=5,
        ),
        "title": ParagraphStyle(
            "ClosureTitle",
            parent=sample["Title"],
            fontName="APClosureArial-Bold",
            fontSize=23,
            leading=27,
            alignment=TA_LEFT,
            textColor=PDF_DARK_BLUE,
            spaceAfter=7,
        ),
        "subtitle": ParagraphStyle(
            "ClosureSubtitle",
            parent=sample["BodyText"],
            fontName="APClosureArial",
            fontSize=12,
            leading=15,
            textColor=PDF_DARK_GRAY,
            spaceAfter=14,
        ),
        "h1": ParagraphStyle(
            "ClosureH1",
            parent=sample["Heading1"],
            fontName="APClosureArial-Bold",
            fontSize=15.5,
            leading=19,
            textColor=PDF_DARK_BLUE,
            spaceBefore=2,
            spaceAfter=8,
        ),
        "caption": ParagraphStyle(
            "ClosureCaption",
            parent=sample["BodyText"],
            fontName="APClosureArial-Bold",
            fontSize=8.2,
            leading=10,
            alignment=TA_CENTER,
            textColor=PDF_DARK_BLUE,
            spaceAfter=4,
        ),
        "metric": ParagraphStyle(
            "ClosureMetric",
            parent=sample["BodyText"],
            fontName="APClosureArial-Bold",
            fontSize=14.5,
            leading=17,
            alignment=TA_CENTER,
            textColor=PDF_DARK_BLUE,
        ),
        "metric_label": ParagraphStyle(
            "ClosureMetricLabel",
            parent=sample["BodyText"],
            fontName="APClosureArial",
            fontSize=7,
            leading=8.8,
            alignment=TA_CENTER,
            textColor=PDF_DARK_GRAY,
        ),
        "callout": ParagraphStyle(
            "ClosureCallout",
            parent=sample["BodyText"],
            fontName="APClosureArial-Bold",
            fontSize=9.2,
            leading=12.3,
            textColor=PDF_GREEN,
        ),
    }


def pdf_paragraph(value: object, style: ParagraphStyle) -> Paragraph:
    escaped = html.escape(str(value)).replace("\n", "<br/>")
    return Paragraph(escaped, style)


def pdf_table(
    rows: Sequence[Sequence[object]],
    widths: Sequence[float],
    styles: dict[str, ParagraphStyle],
    *,
    header: bool = True,
) -> LongTable:
    rendered = []
    for row_index, row in enumerate(rows):
        rendered.append(
            [
                value
                if isinstance(value, (Paragraph, Image, Table, Spacer))
                else pdf_paragraph(
                    value,
                    styles["table_header" if header and row_index == 0 else "table"],
                )
                for value in row
            ]
        )
    table = LongTable(
        rendered,
        colWidths=list(widths),
        repeatRows=1 if header else 0,
        hAlign="LEFT",
    )
    commands = [
        ("GRID", (0, 0), (-1, -1), 0.45, PDF_MID_GRAY),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    if header:
        commands.append(("BACKGROUND", (0, 0), (-1, 0), PDF_LIGHT_GRAY))
    table.setStyle(TableStyle(commands))
    return table


def pdf_callout(text: str, styles: dict[str, ParagraphStyle]) -> Table:
    table = Table([[pdf_paragraph(text, styles["callout"])]], colWidths=[7.06 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), PDF_PALE_GREEN),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#A7D5B7")),
                ("LEFTPADDING", (0, 0), (-1, -1), 9),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    return table


def pdf_metrics(styles: dict[str, ParagraphStyle]) -> Table:
    metrics = (
        ("30/30", "сайтов в матрице"),
        ("60/60", "обработчиков приняли заявки"),
        ("60/60", "просмотров: компьютер/телефон"),
        ("30/30", "маршрутов получателей"),
    )
    cells = [
        [
            pdf_paragraph(value, styles["metric"]),
            pdf_paragraph(label, styles["metric_label"]),
        ]
        for value, label in metrics
    ]
    table = Table([cells], colWidths=[7.06 * inch / 4] * 4)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), PDF_PALE_BLUE),
                ("GRID", (0, 0), (-1, -1), 0.55, colors.HexColor("#B8D4E8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    return table


def pdf_image(path: Path, max_width: float, max_height: float) -> Image:
    with PILImage.open(path) as source:
        width, height = source.size
    scale = min(max_width / width, max_height / height)
    return Image(str(path), width=width * scale, height=height * scale)


def draw_pdf_footer(canvas, document) -> None:
    canvas.saveState()
    canvas.setFont("APClosureArial", 7.2)
    canvas.setFillColor(PDF_DARK_GRAY)
    canvas.drawString(0.72 * inch, 0.34 * inch, "ГК «АП-Риал» | Итоговая проверка форм")
    canvas.drawRightString(7.78 * inch, 0.34 * inch, f"Страница {document.page}")
    canvas.restoreState()


def build_pdf(evidence: dict[str, Any]) -> None:
    styles = pdf_styles()
    story: list[object] = []

    story.extend(
        [
            pdf_paragraph("ИТОГОВЫЙ ОТЧЁТ", styles["kicker"]),
            pdf_paragraph("Исправление и повторная проверка форм", styles["title"]),
            pdf_paragraph("Спорный пакет работ по 30 сайтам ГК «АП-Риал»", styles["subtitle"]),
            pdf_callout(
                "Повторная проверка завершена по полной матрице. На 30 включённых сайтах несоответствий не осталось.",
                styles,
            ),
            Spacer(1, 10),
            pdf_metrics(styles),
            Spacer(1, 13),
            pdf_paragraph(f"Дата проверки: {REPORT_DATE}", styles["small"]),
            pdf_paragraph(
                "Объём отчёта: только последние исправления форм, почтовых маршрутов и повторная приёмка.",
                styles["small"],
            ),
            PageBreak(),
            pdf_paragraph(CLIENT_SECTION_TITLES[0], styles["h1"]),
            pdf_paragraph(
                "Замечания были развернуты на весь список однотипных сайтов, а не только на отдельные примеры. Ниже приведены критерии, по которым выполнена повторная приёмка.",
                styles["body"],
            ),
        ]
    )
    requirement_rows = [["Пункт", "Критерий", "Статус"]] + [
        [title, criterion, "Готово"] for title, criterion in REQUIREMENTS
    ]
    requirement_table = pdf_table(
        requirement_rows,
        [1.62 * inch, 4.64 * inch, 0.8 * inch],
        styles,
    )
    requirement_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (2, 1), (2, -1), PDF_PALE_GREEN),
                ("TEXTCOLOR", (2, 1), (2, -1), PDF_GREEN),
            ]
        )
    )
    story.extend(
        [
            requirement_table,
            Spacer(1, 5),
            pdf_paragraph(
                "Не входили в этот пакет по исходному указанию клиента: "
                + ", ".join(EXCLUDED_SITES)
                + ".",
                styles["small"],
            ),
            pdf_paragraph(CLIENT_SECTION_TITLES[1], styles["h1"]),
        ]
    )
    completed = (
        "Единые требования к двум формам применены ко всем 30 сайтам из матрицы.",
        "Исправлена логика отправки: подтверждение показывается только после ответа обработчика.",
        "Почтовые получатели сверены с исходной таблицей; изменённые маршруты возвращены на штатные адреса.",
        "После изменения маршрутов обе формы на затронутых сайтах отправлены повторно.",
        "Каждый сайт заново открыт и просмотрен на компьютере и телефоне; обе формы проверены отдельно.",
    )
    story.extend(pdf_paragraph(f"• {item}", styles["body"]) for item in completed)
    story.extend(
        [
            PageBreak(),
            pdf_paragraph(CLIENT_SECTION_TITLES[2], styles["h1"]),
            pdf_metrics(styles),
            Spacer(1, 7),
        ]
    )
    results = (
        "60 из 60 обработчиков приняли валидную контрольную заявку и вернули согласованное подтверждение.",
        "18 из 18 форм на сайтах с изменённым маршрутом повторно приняли заявку после исправления адресов.",
        "30 из 30 актуальных конфигураций отправки совпадают с адресами из клиентской матрицы.",
        "16 дополнительных форм Contact Form 7 на семи сайтах отдельно проверены: получатели указаны корректно.",
        "Почтовые домены всех 30 получателей имеют рабочие MX-записи.",
        "Для сквозного контроля в доступном целевом ящике info@medlic.spb.ru подтверждены оба тестовых письма.",
        "60 из 60 контрольных просмотров страниц (30 сайтов × компьютер/телефон) завершены без замечаний; проверено 120 окон форм.",
    )
    story.extend(pdf_paragraph(f"• {item}", styles["body"]) for item in results)
    story.extend(
        [
            pdf_callout(
                "Итог: по включённому списку не осталось незакрытых пунктов. Дополнительные данные от клиента для этого пакета не требуются.",
                styles,
            ),
            PageBreak(),
            pdf_paragraph(CLIENT_SECTION_TITLES[3], styles["h1"]),
            pdf_paragraph(
                "В таблице указан фактический получатель, прочитанный из рабочей конфигурации каждого сайта после исправлений.",
                styles["body"],
            ),
        ]
    )
    matrix_rows = [["№", "Сайт", "Получатель", "Результат"]]
    matrix_by_domain = {
        item["domain"]: item for item in evidence["recipient_matrix"]["sites"]
    }
    for index, domain in enumerate(INCLUDED_SITES, start=1):
        item = matrix_by_domain[domain]
        matrix_rows.append(
            [
                index,
                f"https://{domain}/",
                item["expected_recipient"],
                "Обе формы проверены",
            ]
        )
    matrix_table = pdf_table(
        matrix_rows,
        [0.38 * inch, 1.82 * inch, 2.2 * inch, 2.66 * inch],
        styles,
    )
    matrix_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (3, 1), (3, -1), PDF_PALE_GREEN),
                ("TEXTCOLOR", (3, 1), (3, -1), PDF_GREEN),
            ]
        )
    )
    story.extend(
        [
            matrix_table,
            PageBreak(),
            pdf_paragraph(CLIENT_SECTION_TITLES[4], styles["h1"]),
            pdf_paragraph(
                "Каждый лист содержит три сайта. Для каждого сайта показаны обе формы в компьютерной и мобильной версиях. Снимки сделаны после окончательной настройки почтовых маршрутов.",
                styles["body"],
            ),
        ]
    )
    for index, sheet in enumerate(evidence["sheets"], start=1):
        if index > 1:
            story.append(PageBreak())
        story.append(pdf_paragraph(f"Лист {index} из 10", styles["caption"]))
        max_height = 7.55 * inch if index == 1 else 8.35 * inch
        story.append(pdf_image(sheet, 7.06 * inch, max_height))

    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=LETTER,
        leftMargin=0.72 * inch,
        rightMargin=0.72 * inch,
        topMargin=0.68 * inch,
        bottomMargin=0.58 * inch,
        title="Исправление и повторная проверка форм ГК АП-Риал",
        author="Никита",
    )
    document.build(story, onFirstPage=draw_pdf_footer, onLaterPages=draw_pdf_footer)


def validate_document_outputs() -> dict[str, Any]:
    with zipfile.ZipFile(OUTPUT_DOCX) as archive:
        broken_member = archive.testzip()
        document_xml = archive.read("word/document.xml").decode("utf-8")
    if broken_member is not None:
        raise ValueError(f"DOCX archive contains a broken member: {broken_member}")

    document = Document(OUTPUT_DOCX)
    docx_text_parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            docx_text_parts.extend(cell.text for cell in row.cells)
    docx_text = "\n".join(docx_text_parts)
    external_links = sum(
        1 for relationship in document.part.rels.values() if relationship.is_external
    )
    page_breaks = document_xml.count('w:type="page"')

    pdf_reader = PdfReader(str(OUTPUT_PDF))
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
    required_text = (
        *CLIENT_SECTION_TITLES,
        *INCLUDED_SITES,
        "60 из 60 обработчиков",
        "30 из 30 актуальных конфигураций",
    )
    missing_docx = [value for value in required_text if value not in docx_text]
    missing_pdf = [value for value in required_text if value not in pdf_text]
    forbidden = ("видеофон", "оплата", "счёт №", "агентские сценарии")
    forbidden_docx = [value for value in forbidden if value.lower() in docx_text.lower()]
    forbidden_pdf = [value for value in forbidden if value.lower() in pdf_text.lower()]

    checks = {
        "docx_zip_ok": broken_member is None,
        "docx_tables": len(document.tables),
        "docx_inline_images": len(document.inline_shapes),
        "docx_external_links": external_links,
        "docx_explicit_page_breaks": page_breaks,
        "docx_missing_required_text": missing_docx,
        "docx_forbidden_text": forbidden_docx,
        "pdf_pages": len(pdf_reader.pages),
        "pdf_missing_required_text": missing_pdf,
        "pdf_forbidden_text": forbidden_pdf,
    }
    if len(document.inline_shapes) != 10:
        raise ValueError(f"Expected 10 DOCX evidence images, got {len(document.inline_shapes)}")
    if external_links != 30:
        raise ValueError(f"Expected 30 DOCX site links, got {external_links}")
    if page_breaks != 13:
        raise ValueError(f"Expected 13 explicit DOCX page breaks, got {page_breaks}")
    if len(pdf_reader.pages) != 14:
        raise ValueError(f"Expected 14 PDF pages, got {len(pdf_reader.pages)}")
    if missing_docx or missing_pdf or forbidden_docx or forbidden_pdf:
        raise ValueError(f"Document content validation failed: {checks}")
    return checks


def build_audit(evidence: dict[str, Any]) -> dict[str, Any]:
    source_paths = (
        VISUAL_QA,
        RECIPIENT_MATRIX,
        DELIVERY_QA,
        ROUTE_RETEST,
        CF7_VERIFY,
        MX_AUDIT,
        MAILBOX_PROOF,
        MAILBOX_SCREENSHOT,
        *evidence["sheets"],
    )
    document_qa = validate_document_outputs()
    audit = {
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "scope": {
            "included_sites": list(INCLUDED_SITES),
            "excluded_sites": list(EXCLUDED_SITES),
            "requirements": [
                {"title": title, "criterion": criterion, "status": "passed"}
                for title, criterion in REQUIREMENTS
            ],
        },
        "verification": {
            "visual": evidence["visual"],
            "functional_submissions": evidence["functional_submissions"],
            "route_retest_submissions": evidence["route_retest_submissions"],
            "recipient_routes": evidence["recipient_matrix"]["summary"],
            "legacy_cf7_forms": evidence["cf7_forms"],
            "mx_sites": evidence["mx_sites"],
            "mailbox_delivery_example": evidence["mailbox"],
            "document_qa": document_qa,
        },
        "sources": [
            {
                "path": str(path.relative_to(ROOT)).replace("\\", "/"),
                "sha256": sha256_file(path),
            }
            for path in source_paths
        ],
        "outputs": {
            "docx": {
                "path": str(OUTPUT_DOCX.relative_to(ROOT)).replace("\\", "/"),
                "sha256": sha256_file(OUTPUT_DOCX),
            },
            "pdf": {
                "path": str(OUTPUT_PDF.relative_to(ROOT)).replace("\\", "/"),
                "sha256": sha256_file(OUTPUT_PDF),
                "pages": len(PdfReader(str(OUTPUT_PDF)).pages),
            },
            "email": {
                "path": str(OUTPUT_EMAIL.relative_to(ROOT)).replace("\\", "/"),
                "sha256": sha256_file(OUTPUT_EMAIL),
            },
        },
        "client_contacted": False,
    }
    OUTPUT_AUDIT.write_text(
        json.dumps(audit, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return audit


def main() -> None:
    evidence = validate_sources()
    OUTPUT_EMAIL.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_EMAIL.write_text(build_client_email_text(), encoding="utf-8")
    build_docx(evidence)
    build_pdf(evidence)
    audit = build_audit(evidence)
    print(json.dumps(audit, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
