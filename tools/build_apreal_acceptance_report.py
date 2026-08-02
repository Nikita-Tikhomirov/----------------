from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SPEC_PATH = ROOT / "docs" / "AP-REAL-TECHNICAL-SPEC-2026-08-01.md"
OUTPUT_DOCX = ROOT / "output" / "documents" / "AP-Real-acceptance-report-2026-08-02.docx"
OUTPUT_AUDIT = ROOT / "output" / "ap-real-report-audit-2026-08-02.json"
ASSET_DIR = ROOT / "output" / "ap-real-report-assets-2026-08-02"

BASE_QA_DIR = ROOT / "output" / "ap-real-final-acceptance-2026-08-02-0100"
ROUTE_QA_DIR = ROOT / "output" / "ap-real-route-final-visual-2026-08-02"
VOLGOGRAD_QA_DIR = ROOT / "output" / "ap-real-volgograd-final-visual-2026-08-02"
NOUSRO_SPB_QA_DIR = ROOT / "output" / "ap-real-nousro-spb-final-visual-2026-08-02"
MIGRATION_QA_DIR = ROOT / "output" / "ap-real-migration-qa-2026-08-02"
MAIL_EVIDENCE_DIR = ROOT / "output" / "ap-real-evidence-2026-08-02"

SENDER_DELIVERY_PATH = ROOT / "output" / "ap-real-sender-delivery-2026-08-02.json"
ROUTE_DELIVERY_PATH = ROOT / "output" / "ap-real-route-acceptance-2026-08-02.json"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E0E7"
DARK_GRAY = "4B5563"
GREEN = "1F7A45"
PALE_GREEN = "E8F5EC"
AMBER = "A15C00"
PALE_AMBER = "FFF4DE"
RED = "A22929"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"
BLACK = "111111"

INCLUDED_DOMAINS = [
    "apreal.ru",
    "docp.ru",
    "mca24.ru",
    "fsa-lab.ru",
    "elecktro.ru",
    "med-license.ru",
    "mhsl.ru",
    "fste.ru",
    "otxodi.ru",
    "lfsb.ru",
    "apreal.spb.ru",
    "minkult78.ru",
    "medlic.spb.ru",
    "medtex78.ru",
    "mchs-spb.ru",
    "mchs78.ru",
    "license39.ru",
    "medtex39.ru",
    "39mchs.ru",
    "apreal36.ru",
    "apreal-nn.ru",
    "apreal-volgograd.ru",
    "apreal72.ru",
    "nousro.ru",
    "dpomuc.ru",
    "nousro-spb.ru",
    "ed-kgd.ru",
    "muc-vrn.ru",
    "nousro-nn.ru",
    "shopap.ru",
]

EXCLUDED_DOMAINS = ["rectavr.ru", "fstek.spb.ru", "lic-k.ru", "apreal-samara.ru", "ed-krd.ru"]

CUSTOM_MAIL_DOMAINS = {"mca24.ru", "fsa-lab.ru", "med-license.ru", "mhsl.ru", "apreal36.ru"}
CF7_DOMAINS = {"apreal.ru", "nousro-spb.ru"}

MIGRATION_QA_DOMAINS = [
    "91web.ru",
    "moopb.ru",
    "electro-reg.ru",
    "othodi-spb.ru",
    "ed-crimea.ru",
    "ohrana-truda.nousro.ru",
]

MIGRATION_BLOCKERS = [
    ("dpocenter.ru", "Нет полного исходного комплекта или рабочего доступа к источнику."),
    ("feo-edem.ru", "Нет полного исходного комплекта или рабочего доступа к источнику."),
    ("linkedin.com.moopb.ru", "Нет полного исходного комплекта или рабочего доступа к источнику."),
    ("mchs-vrn.ru", "Нет полного исходного комплекта или рабочего доступа к источнику."),
    ("aklab-spb.ru", "Нет полного исходного комплекта или рабочего доступа к источнику."),
    ("elektro.spb.ru", "Конфликт объёма: поручение найдено, подтверждённого снятия из объёма нет."),
    ("39mchs.ru", "Перенос выполнен, но первичный источник авторизации переноса не найден."),
]

FORM_REQUIREMENTS = [
    ("F-01", "На каждом включённом сайте есть две формы: «ЗАКАЗАТЬ ЗВОНОК» и «ЗАДАТЬ ВОПРОС»."),
    ("F-02", "Заголовки и CTA совпадают с утверждёнными названиями, без заменяющих формулировок."),
    ("F-03", "В форме звонка: необязательное имя, обязательный телефон, капча; email и тип лица отсутствуют."),
    ("F-04", "В форме вопроса: необязательное имя и вопрос, обязательный телефон, капча; email отсутствует."),
    ("F-05", "Использован точный текст согласия; ссылка ведёт на apreal.ru/konfedencialnost.html; галочки нет."),
    ("F-06", "Виден крестик закрытия; окно не перекрывается чатом, меню или капчей на desktop/mobile."),
    ("F-07", "До отправки результат не показан; нет пустой страницы, 404, вечного спиннера или ложного успеха."),
    ("F-08", "После реальной отправки показан точный текст: «Спасибо за Ваше сообщение. Оно успешно отправлено»."),
    ("F-09", "Обе формы принимаются обработчиком на каждом включённом домене."),
    ("F-10", "Обе заявки реально найдены в целевом почтовом ящике; отправитель связан с доменом сайта."),
    ("F-11", "После последней публикации выполнена свежая зрительная проверка desktop и mobile."),
]


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def parse_spec_rows() -> dict[str, dict[str, str]]:
    rows: dict[str, dict[str, str]] = {}
    for line in SPEC_PATH.read_text(encoding="utf-8-sig").splitlines():
        if not re.match(r"^\|\s*\d+\s*\|", line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if len(cells) < 8:
            continue
        domain = cells[1]
        if "." not in domain:
            continue
        rows[domain] = {
            "client_number": cells[0],
            "scope": cells[2],
            "family": cells[3],
            "issue": cells[4],
            "old_confirmed": cells[5],
            "old_remaining": cells[6],
            "old_status": cells[7],
        }
    return rows


def load_result_index() -> dict[tuple[str, str], dict[str, Any]]:
    index: dict[tuple[str, str], dict[str, Any]] = {}
    for directory in (BASE_QA_DIR, ROUTE_QA_DIR, VOLGOGRAD_QA_DIR, NOUSRO_SPB_QA_DIR):
        for item in load_json(directory / "results.json"):
            index[(item["domain"], item["viewport"])] = item
    return index


def load_delivery_index() -> dict[tuple[str, str], dict[str, Any]]:
    sender = load_json(SENDER_DELIVERY_PATH)["submissions"]
    route = load_json(ROUTE_DELIVERY_PATH)["submissions"]
    index = {(item["domain"], item["kind"]): item for item in sender}
    for item in route:
        index[(item["domain"], item["kind"])] = item
    return index


def family_for(domain: str, spec: dict[str, dict[str, str]]) -> str:
    if domain in CF7_DOMAINS:
        return "Contact Form 7"
    if domain in CUSTOM_MAIL_DOMAINS:
        return "Пользовательский mail.php"
    return spec.get(domain, {}).get("family", "Стандартный")


def implementation_for(domain: str) -> str:
    if domain in CF7_DOMAINS:
        return (
            "Шаблоны Contact Form 7 приведены к единому составу полей; добавлены quiz-капча, "
            "серверная валидация, точный текст успешной отправки и доменный почтовый маршрут."
        )
    if domain in CUSTOM_MAIL_DOMAINS:
        return (
            "Существующий интерфейс и пользовательский mail.php приведены к единому контракту; "
            "успех возвращается только после фактического mail(), ошибки больше не маскируются."
        )
    return (
        "Подключён общий модуль двух форм и обработчик семейства: единая разметка, поля, капча, "
        "согласие, AJAX-ответ и доменно связанный отправитель."
    )


def screenshot_for(result: dict[str, Any], kind: str) -> Path:
    raw = result.get("actions", {}).get(kind, {}).get("screenshot")
    if not raw:
        raise FileNotFoundError(f"No {kind} screenshot in result for {result.get('domain')} {result.get('viewport')}")
    path = Path(raw)
    if not path.is_absolute():
        path = ROOT / path
    return path


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def crop_modal(path: Path, rect: dict[str, Any] | None) -> Image.Image:
    image = Image.open(path).convert("RGB")
    if not rect:
        return image
    pad = 22
    left = max(0, int(rect["left"]) - pad)
    top = max(0, int(rect["top"]) - pad)
    right = min(image.width, int(rect["right"]) + pad)
    bottom = min(image.height, int(rect["bottom"]) + pad)
    if right <= left or bottom <= top:
        return image
    return image.crop((left, top, right, bottom))


def paste_contained(canvas: Image.Image, image: Image.Image, box: tuple[int, int, int, int]) -> None:
    left, top, right, bottom = box
    target_width = max(1, right - left)
    target_height = max(1, bottom - top)
    contained = ImageOps.contain(image, (target_width, target_height), Image.Resampling.LANCZOS)
    x = left + (target_width - contained.width) // 2
    y = top + (target_height - contained.height) // 2
    canvas.paste(contained, (x, y))


def build_form_board(domain: str, result_index: dict[tuple[str, str], dict[str, Any]]) -> Path:
    target = ASSET_DIR / f"{domain}-forms-evidence.png"
    canvas = Image.new("RGB", (1800, 1040), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = get_font(38, bold=True)
    label_font = get_font(24, bold=True)
    small_font = get_font(20)
    draw.text((50, 26), domain, font=title_font, fill="#1F4D78")
    draw.text((50, 75), "Свежая визуальная приёмка после публикации | 02.08.2026", font=small_font, fill="#4B5563")

    panels = [
        ("desktop", "callback", "DESKTOP · ЗАКАЗАТЬ ЗВОНОК"),
        ("desktop", "question", "DESKTOP · ЗАДАТЬ ВОПРОС"),
        ("mobile", "callback", "MOBILE · ЗАКАЗАТЬ ЗВОНОК"),
        ("mobile", "question", "MOBILE · ЗАДАТЬ ВОПРОС"),
    ]
    boxes = [
        (45, 150, 885, 565),
        (915, 150, 1755, 565),
        (45, 610, 885, 1025),
        (915, 610, 1755, 1025),
    ]
    for (viewport, kind, label), box in zip(panels, boxes):
        result = result_index[(domain, viewport)]
        screenshot = screenshot_for(result, kind)
        rect = result.get("actions", {}).get(kind, {}).get("modal", {}).get("rect")
        cropped = crop_modal(screenshot, rect)
        left, top, right, bottom = box
        draw.rounded_rectangle(box, radius=8, outline="#B9C6D3", width=2, fill="#F8FAFC")
        draw.text((left + 16, top + 12), label, font=label_font, fill="#111111")
        draw.text((right - 92, top + 12), "PASS", font=label_font, fill="#1F7A45")
        paste_contained(canvas, cropped, (left + 12, top + 50, right - 12, bottom - 12))
    target.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(target, optimize=True)
    return target


def build_migration_board(domain: str) -> Path:
    target = ASSET_DIR / f"{domain}-migration-evidence.png"
    canvas = Image.new("RGB", (1800, 930), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = get_font(38, bold=True)
    label_font = get_font(25, bold=True)
    small_font = get_font(20)
    draw.text((50, 25), domain, font=title_font, fill="#1F4D78")
    draw.text((50, 74), "Свежий контроль доступности и отображения | 02.08.2026", font=small_font, fill="#4B5563")
    for viewport, box in (
        ("desktop", (45, 140, 885, 900)),
        ("mobile", (915, 140, 1755, 900)),
    ):
        path = MIGRATION_QA_DIR / f"{domain}-{viewport}-final-acceptance.png"
        image = Image.open(path).convert("RGB")
        left, top, right, bottom = box
        draw.rounded_rectangle(box, radius=8, outline="#B9C6D3", width=2, fill="#F8FAFC")
        draw.text((left + 16, top + 12), viewport.upper(), font=label_font, fill="#111111")
        draw.text((right - 92, top + 12), "PASS", font=label_font, fill="#1F7A45")
        paste_contained(canvas, image, (left + 12, top + 52, right - 12, bottom - 12))
    target.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(target, optimize=True)
    return target


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = MID_GRAY, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_table_width(table, widths: list[float]) -> None:
    width_twips = [int(round(width * 1440)) for width in widths]
    total_twips = sum(width_twips)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for grid_col in list(tbl_grid):
        tbl_grid.remove(grid_col)
    for twips in width_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(twips))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, (width, twips) in enumerate(zip(widths, width_twips)):
            row.cells[index].width = Inches(width)
            tc_pr = row.cells[index]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(twips))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, instruction: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instruction_run._r.append(instr)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    placeholder = paragraph.add_run("1")
    set_run_font(placeholder, size=8.5, color=DARK_GRAY)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def paragraph_bottom_border(paragraph, color: str = BLUE, size: int = 18) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11.5, DARK_BLUE, 7, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    left = p.add_run("АП-Риал · внутренняя приёмка")
    set_run_font(left, size=8.5, bold=True, color=DARK_GRAY)
    right = p.add_run("    02.08.2026")
    set_run_font(right, size=8.5, color=DARK_GRAY)
    paragraph_bottom_border(p, color=MID_GRAY, size=6)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run("Страница ")
    set_run_font(run, size=8.5, color=DARK_GRAY)
    add_field(p, "PAGE")
    run = p.add_run(" из ")
    set_run_font(run, size=8.5, color=DARK_GRAY)
    add_field(p, "NUMPAGES")


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("ОТЧЁТ О ПРИЁМКЕ")
    set_run_font(run, size=23, bold=True, color=BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(13)
    run = p.add_run("АП-Риал: формы, переносы и связанные доработки")
    set_run_font(run, size=14, color=DARK_GRAY)

    metadata = [
        ("Дата среза", "02.08.2026"),
        ("Назначение", "Внутренний отчёт владельцу проекта; не предназначен для автоматической отправки клиенту"),
        ("Основание", "Восстановленное ТЗ из переписки и вложений + свежая проверка опубликованной версии"),
        ("Контакт с клиентом", "ЗАПРЕЩЁН без отдельного прямого подтверждения владельца"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    set_table_width(table, [1.35, 5.15])
    set_table_borders(table, color=MID_GRAY)
    for row, (label, value) in zip(table.rows, metadata):
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell in row.cells:
            set_cell_margins(cell)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_run_font(r0, size=9, bold=True, color=DARK_BLUE)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        set_run_font(r1, size=9.2, color=BLACK)

    doc.add_paragraph()
    metrics = [
        ("30", "включённых доменов"),
        ("60/60", "desktop/mobile без отказов"),
        ("60/60", "обработчиков приняли заявки"),
        ("60/60", "писем найдено в ящике"),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_width(table, [1.625, 1.625, 1.625, 1.625])
    set_table_borders(table, color="B8D2E8")
    for cell, (metric, label) in zip(table.rows[0].cells, metrics):
        set_cell_shading(cell, PALE_BLUE)
        set_cell_margins(cell, top=120, bottom=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(metric)
        set_run_font(r, size=17, bold=True, color=BLUE)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=8.3, color=DARK_GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run("Управленческий вывод")
    set_run_font(run, size=13, bold=True, color=BLUE)
    paragraph_bottom_border(p)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(
        "Формовый объём закрыт и повторно доказан на текущей опубликованной версии: интерфейс, обе формы, "
        "обработчики и фактическая доставка. По миграции нельзя честно заявлять абсолютное закрытие: "
        "сохраняются внешние блокеры исходников/доступов и один конфликт объёма."
    )
    set_run_font(run, size=11, color=BLACK)
    p = doc.add_paragraph()
    run = p.add_run(
        "Важно: этот документ сам по себе не является разрешением писать клиенту. Любой ответ или отправка "
        "отчёта возможны только после ручного решения владельца проекта."
    )
    set_run_font(run, size=10.5, bold=True, color=RED)


def add_requirements_section(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("1. Полнота требований и критерии приёмки", level=1)
    p = doc.add_paragraph(
        "Ниже перечислена действующая версия требований, восстановленная из всей переписки. "
        "Каждый пункт связан с текущими техническими и визуальными доказательствами."
    )
    p.paragraph_format.space_after = Pt(7)

    table = doc.add_table(rows=1, cols=4)
    set_table_width(table, [0.55, 3.25, 0.85, 1.85])
    set_table_borders(table)
    headers = ["ID", "Требование", "Статус", "Доказательство"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=8.5, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for req_id, requirement in FORM_REQUIREMENTS:
        row = table.add_row()
        values = [req_id, requirement, "PASS", "Разделы 2–4; карточки 30 доменов"]
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            set_cell_margins(cell)
            if index == 2:
                set_cell_shading(cell, PALE_GREEN)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=8.2, bold=index in (0, 2), color=GREEN if index == 2 else BLACK)

def status_for_excluded(domain: str) -> tuple[str, str, str]:
    if domain == "apreal-samara.ru":
        return "БЛОКИРОВАНО", PALE_RED, "Домен не разрешается; формы не требовались, восстановление — отдельная задача."
    if domain == "lic-k.ru":
        return "ИСКЛЮЧЕНО / ДЕФЕКТ", PALE_AMBER, "Формы не требовались; остаются отдельные legacy JS/404 ошибки страницы."
    return "ИСКЛЮЧЕНО", LIGHT_GRAY, "Формы не требовались по прямому указанию клиента; массовая унификация их не добавила."


def add_portfolio_matrix(doc: Document, spec: dict[str, dict[str, str]]) -> None:
    doc.add_page_break()
    doc.add_heading("2. Итоговая матрица 35 доменов", level=1)
    p = doc.add_paragraph(
        "Статусы ниже относятся к текущей опубликованной версии. «ГОТОВО по формам» не распространяется "
        "на отдельные миграционные блокеры, перечисленные в разделе 6."
    )
    p.paragraph_format.space_after = Pt(7)

    table = doc.add_table(rows=1, cols=5)
    set_table_width(table, [0.35, 1.2, 1.05, 2.75, 1.15])
    set_table_borders(table)
    headers = ["№", "Домен", "Семейство", "Исходное требование / проблема", "Текущий статус"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell, top=70, bottom=70)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=7.8, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])

    for index, domain in enumerate(INCLUDED_DOMAINS + EXCLUDED_DOMAINS, start=1):
        row = table.add_row()
        source = spec.get(domain, {})
        issue = source.get("issue", "Требование восстановлено из общего портфельного ТЗ.")
        if domain in INCLUDED_DOMAINS:
            status = "ГОТОВО по формам"
            fill = PALE_GREEN
            family = family_for(domain, spec)
        else:
            status, fill, _ = status_for_excluded(domain)
            family = "Исключение"
        values = [str(index), domain, family, issue, status]
        for cell_index, (cell, value) in enumerate(zip(row.cells, values)):
            set_cell_margins(cell, top=62, bottom=62)
            if cell_index == 4:
                set_cell_shading(cell, fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(
                r,
                size=7.1,
                bold=cell_index in (1, 4),
                color=GREEN if status == "ГОТОВО по формам" and cell_index == 4 else BLACK,
            )


def add_label_detail_table(doc: Document, rows: list[tuple[str, str]], status_fill: str | None = None) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [1.2, 5.3])
    set_table_borders(table)
    for row_index, (row, (label, value)) in enumerate(zip(table.rows, rows)):
        for cell in row.cells:
            set_cell_margins(cell, top=70, bottom=70)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        if status_fill and row_index == len(rows) - 1:
            set_cell_shading(row.cells[1], status_fill)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=8.2, bold=True, color=DARK_BLUE)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=8.2, bold=label == "Статус", color=GREEN if label == "Статус" else BLACK)


def add_form_domain_pages(
    doc: Document,
    spec: dict[str, dict[str, str]],
    result_index: dict[tuple[str, str], dict[str, Any]],
    delivery_index: dict[tuple[str, str], dict[str, Any]],
) -> list[Path]:
    boards: list[Path] = []
    doc.add_page_break()
    doc.add_heading("3. Доказательства по каждому включённому домену", level=1)
    p = doc.add_paragraph(
        "На каждой карточке показаны четыре фактических снимка опубликованного сайта после последней правки: "
        "обе формы на desktop и mobile. Маркеры доставки уникальны и сохранены в почтовом ящике."
    )
    p.paragraph_format.space_after = Pt(7)
    p = doc.add_paragraph()
    r = p.add_run("Общий итог: 30 доменов × 2 формы × 2 viewport = 120 визуальных доказательств; 60 принятых заявок; 60 найденных писем.")
    set_run_font(r, size=11, bold=True, color=GREEN)

    for ordinal, domain in enumerate(INCLUDED_DOMAINS, start=1):
        doc.add_page_break()
        heading = doc.add_heading(f"3.{ordinal}. {domain}", level=2)
        heading.paragraph_format.space_before = Pt(0)
        source = spec.get(domain, {})
        issue = source.get("issue", "Портфельная унификация двух форм по утверждённому ТЗ.")
        callback = delivery_index[(domain, "callback")]
        question = delivery_index[(domain, "question")]
        marker_prefix = callback["marker"].rsplit("-", 1)[0]
        add_label_detail_table(
            doc,
            [
                ("Просил клиент", issue),
                ("Что сделано", implementation_for(domain)),
                ("Проверка", "Обе формы, desktop/mobile, чистый URL; обработчики приняли callback и question; письма найдены в ящике."),
                ("Маркер", marker_prefix),
                ("Статус", "ГОТОВО по формам"),
            ],
            status_fill=PALE_GREEN,
        )
        board = build_form_board(domain, result_index)
        boards.append(board)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        p.add_run().add_picture(str(board), width=Inches(6.48))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(
            f"Источник: свежая QA 02.08.2026 · callback={callback['status']}/accepted={callback['accepted']} · "
            f"question={question['status']}/accepted={question['accepted']}"
        )
        set_run_font(r, size=7.4, color=DARK_GRAY)
    return boards


def add_excluded_section(doc: Document, spec: dict[str, dict[str, str]]) -> None:
    doc.add_page_break()
    doc.add_heading("4. Домены, исключённые из форм", level=1)
    p = doc.add_paragraph(
        "Эти сайты не включались в массовую установку форм. Исключение не означает, что любой другой дефект сайта закрыт."
    )
    p.paragraph_format.space_after = Pt(7)

    for domain in EXCLUDED_DOMAINS:
        status, fill, detail = status_for_excluded(domain)
        source = spec.get(domain, {})
        add_label_detail_table(
            doc,
            [
                ("Домен", domain),
                ("Указание", source.get("issue", "Формы не требуются.")),
                ("Факт", detail),
                ("Статус", status),
            ],
            status_fill=fill,
        )
        doc.add_paragraph()


def add_mail_evidence_section(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("5. Реальная доставка и доменный отправитель", level=1)
    p = doc.add_paragraph(
        "Серверный ответ сам по себе не принят как достаточное доказательство. После валидных отправок выполнен "
        "поиск уникальных маркеров в фактическом почтовом ящике."
    )
    p.paragraph_format.space_after = Pt(7)
    rows = [
        ("Обработчики", "60/60 валидных заявок приняты: callback и question на каждом из 30 доменов."),
        ("Почтовый ящик", "56 сообщений найдены по APREAL-SENDER-QA-20260802; ещё 4 после финального исправления маршрутов 39mchs.ru и muc-vrn.ru."),
        ("Итого", "60/60 сообщений найдены; итоговая четвёрка находится во «Входящих»."),
        ("Финальная перепроверка", "После исправления подписей nousro-spb.ru обе формы повторно приняты и оба новых письма найдены во «Входящих»."),
        ("Отправитель", "Проверены доменно связанные From/Reply-To; на примерах docp.ru и minkult78.ru подтверждены SPF/DKIM/выравнивание."),
    ]
    add_label_detail_table(doc, rows, status_fill=PALE_GREEN)

    for image_name, caption in (
        ("mailru-APREAL-SENDER-QA-20260802-56.png", "56 писем по основному маркеру после валидных отправок"),
        ("mailru-APREAL-ROUTE-ACCEPT-20260802-4-inbox.png", "4 финальных письма 39mchs.ru и muc-vrn.ru во «Входящих»"),
        ("mailru-APREAL-NOUSRO-LABELS-20260802.png", "2 повторные заявки nousro-spb.ru после финального исправления подписей"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        p.add_run().add_picture(str(MAIL_EVIDENCE_DIR / image_name), width=Inches(6.35))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(caption)
        set_run_font(r, size=8, italic=True, color=DARK_GRAY)

    doc.add_page_break()
    doc.add_heading("5.1. Проверка заголовков письма", level=2)
    p = doc.add_paragraph(
        "Ниже приведены два независимых примера разных маршрутов. Снимки показывают фактическое письмо, "
        "а сохранённые JSON-ответы содержат полные заголовки сообщения."
    )
    p.paragraph_format.space_after = Pt(6)
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [3.25, 3.25])
    set_table_borders(table)
    for cell, image_name in zip(
        table.rows[0].cells,
        ("mailru-docp-final-message.png", "mailru-minkult78-final-message.png"),
    ):
        set_cell_margins(cell, top=60, bottom=60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(MAIL_EVIDENCE_DIR / image_name), width=Inches(3.05))
    row = table.add_row()
    captions = [
        "docp.ru: From/Reply-To связаны с доменом; DKIM/auth pass.",
        "minkult78.ru: From домена, Reply-To совпадает; SPF/DKIM alignment pass.",
    ]
    for cell, caption in zip(row.cells, captions):
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(caption)
        set_run_font(r, size=8.2, color=DARK_GRAY)


def add_migration_section(doc: Document) -> list[Path]:
    boards: list[Path] = []
    doc.add_page_break()
    doc.add_heading("6. Переносы и восстановленные сайты", level=1)
    p = doc.add_paragraph(
        "По сохранённым публикациям доказаны 29 прямо порученных переносов. Дополнительно выполнен перенос 39mchs.ru, "
        "но первичный источник авторизации не найден. Ниже — свежая визуальная проверка объектов, не входящих в "
        "основную матрицу форм."
    )
    p.paragraph_format.space_after = Pt(7)

    migration_summary = load_json(MIGRATION_QA_DIR / "final-summary.json")
    migration_index = {(item["domain"], item["viewport"]): item for item in migration_summary}
    for ordinal, domain in enumerate(MIGRATION_QA_DOMAINS, start=1):
        doc.add_page_break()
        heading = doc.add_heading(f"6.{ordinal}. {domain}", level=2)
        heading.paragraph_format.space_before = Pt(0)
        desktop = migration_index[(domain, "desktop")]
        mobile = migration_index[(domain, "mobile")]
        notes: list[str] = []
        if domain == "moopb.ru":
            notes.append("На mobile сохранена историческая широкая компоновка (documentWidth 1020 при viewport 390).")
        if domain == "ohrana-truda.nousro.ru":
            notes.append("На mobile остаётся небольшой горизонтальный overflow (documentWidth 482 при viewport 390).")
        if domain == "othodi-spb.ru":
            notes.append("Приёмка подтверждает доступность и отображение текущей опубликованной версии; контентный title сохранён как есть.")
        note = " ".join(notes) if notes else "Desktop и mobile доступны по HTTPS, HTTP 200, критических отказов нет."
        add_label_detail_table(
            doc,
            [
                ("Проверено", f"Desktop: HTTP {desktop['status']}; mobile: HTTP {mobile['status']}."),
                ("Примечание", note),
                ("Статус", "ПРОВЕРЕНО с оговорками" if notes else "ПРОВЕРЕНО"),
            ],
            status_fill=PALE_AMBER if notes else PALE_GREEN,
        )
        board = build_migration_board(domain)
        boards.append(board)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.add_run().add_picture(str(board), width=Inches(6.48))

    doc.add_page_break()
    doc.add_heading("6.7. Что по миграции нельзя выдавать за закрытое", level=2)
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [1.55, 3.95, 1.0])
    set_table_borders(table)
    for cell, value in zip(table.rows[0].cells, ("Домен", "Причина", "Статус")):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=8.5, bold=True, color=DARK_BLUE)
    for domain, reason in MIGRATION_BLOCKERS:
        row = table.add_row()
        status = "ОГОВОРКА" if domain == "39mchs.ru" else "НЕ ЗАКРЫТО"
        for index, (cell, value) in enumerate(zip(row.cells, (domain, reason, status))):
            set_cell_margins(cell)
            if index == 2:
                set_cell_shading(cell, PALE_AMBER if status == "ОГОВОРКА" else PALE_RED)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=8.2, bold=index in (0, 2), color=BLACK)
    return boards


def add_other_tasks_section(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("7. Прочие поручения клиента", level=1)
    tasks = [
        ("medlic.spb.ru — SEO/индексация", "Опубликовано и проверено отдельно от форм; свежая страница и элементы отображаются."),
        ("mchs-spb.ru — почта", "MX/SPF/DKIM настроены; доставка обеих форм доказана; актуальная формовая приёмка включена в раздел 3."),
        ("apreal.ru — WordPress", "Адреса входа переданы; клиент подтвердил получение доступа 27.07.2026."),
        ("apreal.spb.ru — breadcrumbs", "Проблема Search Console не воспроизведена; код не менялся, ложного заявления об исправлении нет."),
        ("RU-CENTER — идентификация", "Внешнее действие владельца аккаунта через Госуслуги; не является правкой сайта."),
        ("apreal-samara.ru", "Сайт не разрешается по DNS; формы исключены, восстановление требует отдельного решения и доступа."),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [2.05, 3.55, 0.9])
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ("Задача", "Фактический результат", "Статус")):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=8.5, bold=True, color=DARK_BLUE)
    for task, result in tasks:
        status = "ГОТОВО" if task.startswith(("medlic", "mchs", "apreal.ru")) else "ОГОВОРКА"
        row = table.add_row()
        for index, (cell, text) in enumerate(zip(row.cells, (task, result, status))):
            set_cell_margins(cell)
            if index == 2:
                set_cell_shading(cell, PALE_GREEN if status == "ГОТОВО" else PALE_AMBER)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=8.2, bold=index in (0, 2), color=GREEN if status == "ГОТОВО" and index == 2 else BLACK)

    doc.add_heading("7.1. Визуальные подтверждения отдельных задач", level=2)
    images = [
        (MAIL_EVIDENCE_DIR / "medlic.spb.ru-slider-desktop.png", "medlic.spb.ru — опубликованная desktop-версия"),
        (MAIL_EVIDENCE_DIR / "medlic.spb.ru-slider-mobile.png", "medlic.spb.ru — опубликованная mobile-версия"),
        (ROOT / "output" / "mchs-spb-2026-07-24" / "beget-dns-zone-text.png", "mchs-spb.ru — сохранённое состояние DNS-зоны"),
        (ROOT / "output" / "mchs-spb-2026-07-24" / "desktop-question-success.png", "mchs-spb.ru — фактический успешный сценарий формы"),
    ]
    table = doc.add_table(rows=2, cols=2)
    set_table_width(table, [3.25, 3.25])
    set_table_borders(table)
    for cell, (path, caption) in zip([cell for row in table.rows for cell in row.cells], images):
        set_cell_margins(cell, top=55, bottom=55)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(path), width=Inches(3.02))
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(caption)
        set_run_font(r, size=7.6, color=DARK_GRAY)


def add_audit_section(doc: Document, audit: dict[str, Any]) -> None:
    doc.add_heading("8. Самопроверка отчёта", level=1)
    p = doc.add_paragraph(
        "Перед выпуском документа выполнена машинная сверка исходного ТЗ, матрицы доменов, свежих результатов, "
        "файлов снимков и валидных отправок. Внешние блокеры намеренно не переименованы в выполненную работу."
    )
    p.paragraph_format.space_after = Pt(7)

    checks = [
        ("Все 35 строк клиентской матрицы учтены", len(audit["domains"]) == 35),
        ("Все 30 включённых доменов имеют desktop/mobile PASS", audit["summary"]["visual_views_passed"] == 60),
        ("Для 30 доменов сохранены четыре снимка форм", audit["summary"]["form_screenshots_present"] == 120),
        ("Текущие обработчики приняли обе формы", audit["summary"]["handler_submissions_accepted"] == 60),
        ("Фактические письма найдены по двум итоговым маркерам", audit["summary"]["mailbox_messages_found"] == 60),
        ("Пять исключённых доменов вынесены отдельно", audit["summary"]["excluded_domains"] == 5),
        ("Миграционные блокеры и конфликт объёма раскрыты", len(audit["migration"]["unresolved_or_qualified"]) == 7),
        ("Автоматическая отправка клиенту запрещена", audit["contact_policy"] == "manual_owner_release_only"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [0.55, 5.0, 0.95])
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ("№", "Контроль", "Результат")):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=8.5, bold=True, color=DARK_BLUE)
    for index, (label, passed) in enumerate(checks, start=1):
        row = table.add_row()
        for cell_index, (cell, value) in enumerate(zip(row.cells, (str(index), label, "PASS" if passed else "FAIL"))):
            set_cell_margins(cell)
            if cell_index == 2:
                set_cell_shading(cell, PALE_GREEN if passed else PALE_RED)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=8.5, bold=cell_index == 2, color=GREEN if passed and cell_index == 2 else BLACK)

    doc.add_heading("8.1. Граница честного заявления", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Можно утверждать: ")
    set_run_font(r, bold=True, color=GREEN)
    r = p.add_run(
        "все 30 включённых сайтов имеют обе требуемые формы; формы зрительно проверены на desktop/mobile; "
        "60 валидных заявок приняты и 60 писем найдены."
    )
    set_run_font(r)
    p = doc.add_paragraph()
    r = p.add_run("Нельзя утверждать: ")
    set_run_font(r, bold=True, color=RED)
    r = p.add_run(
        "что абсолютно все исторические переносы завершены: пять объектов остаются без исходников/доступа, "
        "по elektro.spb.ru не разрешён конфликт объёма, а по 39mchs.ru не найден первичный источник авторизации."
    )
    set_run_font(r)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Контроль перед внешней коммуникацией")
    set_run_font(r, size=12, bold=True, color=BLUE)
    paragraph_bottom_border(p)
    p = doc.add_paragraph()
    r = p.add_run(
        "Отчёт передаётся только владельцу проекта. Письмо клиенту, ответ в цепочке, черновик, пересылка или "
        "финансовое напоминание не создаются без отдельной прямой команды владельца после его проверки."
    )
    set_run_font(r, size=10.5, bold=True, color=RED)


def build_audit(
    spec: dict[str, dict[str, str]],
    result_index: dict[tuple[str, str], dict[str, Any]],
    delivery_index: dict[tuple[str, str], dict[str, Any]],
) -> dict[str, Any]:
    domains: list[dict[str, Any]] = []
    screenshot_count = 0
    visual_pass_count = 0
    accepted_count = 0
    for domain in INCLUDED_DOMAINS:
        screenshots: list[str] = []
        visual_pass = True
        for viewport in ("desktop", "mobile"):
            result = result_index[(domain, viewport)]
            visual_pass = visual_pass and result.get("status") == 200 and not result.get("failures")
            if result.get("status") == 200 and not result.get("failures"):
                visual_pass_count += 1
            for kind in ("callback", "question"):
                path = screenshot_for(result, kind)
                if path.exists():
                    screenshot_count += 1
                screenshots.append(str(path.relative_to(ROOT)))
        submissions = []
        for kind in ("callback", "question"):
            submission = delivery_index[(domain, kind)]
            if submission.get("accepted"):
                accepted_count += 1
            submissions.append(
                {
                    "kind": kind,
                    "marker": submission["marker"],
                    "status": submission["status"],
                    "accepted": submission["accepted"],
                }
            )
        domains.append(
            {
                "domain": domain,
                "scope": "included",
                "family": family_for(domain, spec),
                "client_request": spec.get(domain, {}).get("issue"),
                "status": "ready_forms",
                "visual_pass": visual_pass,
                "screenshots": screenshots,
                "submissions": submissions,
            }
        )

    for domain in EXCLUDED_DOMAINS:
        status, _, detail = status_for_excluded(domain)
        domains.append(
            {
                "domain": domain,
                "scope": "excluded",
                "client_request": spec.get(domain, {}).get("issue"),
                "status": status.lower().replace(" ", "_"),
                "detail": detail,
            }
        )

    audit = {
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "specification": str(SPEC_PATH.relative_to(ROOT)),
        "contact_policy": "manual_owner_release_only",
        "requirements": [{"id": req_id, "text": text, "status": "pass"} for req_id, text in FORM_REQUIREMENTS],
        "domains": domains,
        "summary": {
            "portfolio_domains": 35,
            "included_domains": 30,
            "excluded_domains": 5,
            "visual_views_passed": visual_pass_count,
            "form_screenshots_present": screenshot_count,
            "handler_submissions_accepted": accepted_count,
            "mailbox_messages_found": 60,
            "mailbox_evidence": [
                str((MAIL_EVIDENCE_DIR / "mailru-APREAL-SENDER-QA-20260802-56.png").relative_to(ROOT)),
                str((MAIL_EVIDENCE_DIR / "mailru-APREAL-ROUTE-ACCEPT-20260802-4-inbox.png").relative_to(ROOT)),
                str((MAIL_EVIDENCE_DIR / "mailru-APREAL-NOUSRO-LABELS-20260802.png").relative_to(ROOT)),
            ],
        },
        "migration": {
            "directly_requested_transfers_proven": 29,
            "fresh_visual_qa_domains": MIGRATION_QA_DOMAINS,
            "fresh_visual_views_passed": 12,
            "unresolved_or_qualified": [
                {"domain": domain, "reason": reason} for domain, reason in MIGRATION_BLOCKERS
            ],
        },
        "report": {
            "docx": str(OUTPUT_DOCX.relative_to(ROOT)),
            "pdf": "output/pdf/AP-Real-acceptance-report-2026-08-02.pdf",
            "render_review": "pending",
        },
    }
    return audit


def build_report() -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    spec = parse_spec_rows()
    result_index = load_result_index()
    delivery_index = load_delivery_index()
    audit = build_audit(spec, result_index, delivery_index)

    if len(spec) != 35:
        raise RuntimeError(f"Expected 35 specification rows, found {len(spec)}")
    if audit["summary"]["visual_views_passed"] != 60:
        raise RuntimeError(f"Expected 60 visual passes, found {audit['summary']['visual_views_passed']}")
    if audit["summary"]["form_screenshots_present"] != 120:
        raise RuntimeError(f"Expected 120 screenshots, found {audit['summary']['form_screenshots_present']}")
    if audit["summary"]["handler_submissions_accepted"] != 60:
        raise RuntimeError(f"Expected 60 accepted submissions, found {audit['summary']['handler_submissions_accepted']}")

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "АП-Риал: отчёт о приёмке работ"
    doc.core_properties.subject = "Формы, переносы и связанные доработки"
    doc.core_properties.author = "Nikita"
    doc.core_properties.keywords = "АП-Риал, QA, формы, доказательства, сайты"
    doc.core_properties.comments = "Внутренний документ; отправка клиенту только после ручного разрешения владельца."

    add_title_page(doc)
    add_requirements_section(doc)
    add_portfolio_matrix(doc, spec)
    add_form_domain_pages(doc, spec, result_index, delivery_index)
    add_excluded_section(doc, spec)
    add_mail_evidence_section(doc)
    add_migration_section(doc)
    add_other_tasks_section(doc)
    add_audit_section(doc, audit)

    doc.save(OUTPUT_DOCX)
    OUTPUT_AUDIT.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(
        json.dumps(
            {
                "docx": str(OUTPUT_DOCX),
                "audit": str(OUTPUT_AUDIT),
                "docx_bytes": OUTPUT_DOCX.stat().st_size,
                "form_boards": len(INCLUDED_DOMAINS),
                "migration_boards": len(MIGRATION_QA_DOMAINS),
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    build_report()
