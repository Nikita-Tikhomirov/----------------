#!/usr/bin/env python3
"""Build the client-facing residual quality report as DOCX, PDF, and JSON."""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as PdfImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib import colors


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output/residual-quality-fixes-20260813"
QA = OUT / "qa"
ASSETS = OUT / "report-assets"
DOCX_PATH = OUT / "Отчет-об-исправлениях-medlic-lfsb-2026-08-13.docx"
PDF_PATH = OUT / "Отчет-об-исправлениях-medlic-lfsb-2026-08-13.pdf"
AUDIT_PATH = OUT / "audit.json"

INK = "17212B"
ACCENT = "7A0B25"
GREEN = "276749"
MUTED = "5B6573"


def crop(name: str, source: Path, box: tuple[int, int, int, int]) -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    target = ASSETS / name
    with Image.open(source) as image:
        image.crop(box).save(target, optimize=True)
    return target


def prepare_assets() -> dict[str, Path]:
    return {
        "lfsb_before": crop(
            "lfsb-before-mobile-top.png",
            OUT / "before/lfsb-mobile-390.png",
            (0, 0, 390, 1000),
        ),
        "lfsb_after": crop(
            "lfsb-after-mobile-top.png",
            QA / "lfsb-home-mobile.png",
            (0, 0, 390, 1000),
        ),
        "medlic_after": crop(
            "medlic-after-mobile-top.png",
            QA / "medlic-home-mobile.png",
            (0, 0, 390, 1250),
        ),
        "fstec_after": crop(
            "lfsb-fstec-dir-after-top.png",
            QA / "lfsb-fstec_dir-mobile.png",
            (0, 0, 390, 1500),
        ),
        "kripto_after": crop(
            "lfsb-kripto-dir-after-top.png",
            QA / "lfsb-kripto_dir-mobile.png",
            (0, 0, 390, 1700),
        ),
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_docx_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Отчёт об исправлениях на medlic.spb.ru и lfsb.ru")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)
    subtitle = doc.add_paragraph("Повторная проверка и устранение обнаруженных недочётов")
    subtitle.style = "Subtitle"
    doc.add_paragraph("13 августа 2026 года")


def add_docx_picture(doc: Document, path: Path, width: float, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = "Caption"


def add_docx_status(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E8F5ED")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN)


def build_docx(assets: dict[str, Path]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size in (("Heading 1", 16), ("Heading 2", 13)):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "АП-Риал | Отчёт о выполненных исправлениях"
    header.style = styles["Caption"]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("13.08.2026")

    add_docx_title(doc)
    add_docx_status(doc, "Все перечисленные ниже недочёты устранены и повторно проверены.")
    doc.add_paragraph(
        "Проведена повторная проверка двух сайтов. Исправлены найденные ошибки в текстах, "
        "мобильной вёрстке и старых подключениях страниц. После публикации сайты заново "
        "проверены на компьютере и телефоне."
    )

    doc.add_heading("1. medlic.spb.ru", level=1)
    doc.add_paragraph(
        "Что было не так: на странице оставались две опечатки, а скрытые пункты мобильного "
        "меню могли увеличивать ширину страницы."
    )
    doc.add_paragraph(
        "Что сделано: «Всеь процессы» заменено на «Все процессы», а «Росздравнадзоррешает» "
        "заменено на «Росздравнадзор решает». Мобильное меню ограничено шириной экрана. "
        "Содержание и внешний вид компьютерной версии не менялись."
    )
    add_docx_picture(doc, assets["medlic_after"], 2.5, "medlic.spb.ru на телефоне после исправления")

    doc.add_page_break()
    doc.add_heading("2. lfsb.ru", level=1)
    doc.add_paragraph(
        "Что было не так: сайт был построен на старой фиксированной вёрстке. На телефоне часть "
        "страниц обрезалась, отдельные карточки накладывались друг на друга, а несколько старых "
        "подключений браузер блокировал как небезопасные или отсутствующие."
    )
    doc.add_paragraph(
        "Что сделано: все 22 публичные страницы адаптированы для экранов телефона; шапка, меню, "
        "текст, изображения, карточки, боковые блоки и подвал теперь располагаются последовательно. "
        "Старые подключения заменены на рабочие локальные или защищённые версии."
    )
    compare = doc.add_table(rows=1, cols=2)
    compare.autofit = False
    for cell, image, caption in (
        (compare.cell(0, 0), assets["lfsb_before"], "До повторной проверки"),
        (compare.cell(0, 1), assets["lfsb_after"], "После исправления"),
    ):
        cell.width = Inches(3.1)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image), width=Inches(2.4))
        cp = cell.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].bold = True

    doc.add_page_break()
    doc.add_heading("3. Длинные страницы и карточки", level=1)
    doc.add_paragraph(
        "Отдельно проверены страницы с большим количеством карточек и старой сложной разметкой. "
        "Карточки выстроены вертикально, текст и изображения больше не перекрывают друг друга."
    )
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    for cell, image, caption in (
        (table.cell(0, 0), assets["fstec_after"], "Лицензия ФСТЭК"),
        (table.cell(0, 1), assets["kripto_after"], "Лицензия на криптографию"),
    ):
        cell.width = Inches(3.1)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image), width=Inches(2.25))
        cp = cell.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].bold = True

    doc.add_page_break()
    doc.add_heading("4. Формы и итоговая проверка", level=1)
    doc.add_paragraph(
        "На обоих сайтах повторно открыты формы «Заказать звонок» и «Задать вопрос» на компьютере "
        "и телефоне. Поля читаемы, имеют одинаковую высоту и не выходят за границы окна."
    )
    form_table = doc.add_table(rows=1, cols=2)
    form_table.autofit = False
    for cell, image, caption in (
        (form_table.cell(0, 0), QA / "lfsb-home-mobile-question.png", "lfsb.ru"),
        (form_table.cell(0, 1), QA / "medlic-home-mobile-question.png", "medlic.spb.ru"),
    ):
        cell.width = Inches(3.1)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image), width=Inches(2.45))
        cp = cell.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].bold = True

    doc.add_heading("Что проверено после публикации", level=2)
    for text in (
        "22 публичные страницы lfsb.ru на экране телефона 390 px;",
        "главная lfsb.ru на 320 px и 1440 px;",
        "главная medlic.spb.ru на 390 px и 1440 px;",
        "открытие обеих форм на каждом сайте;",
        "отсутствие обрезки, перекрытий, ошибок браузера и неудачных загрузок.",
    ):
        doc.add_paragraph(text, style="List Bullet")
    add_docx_status(doc, "Итог: обнаруженные недочёты устранены. Дополнительные данные от заказчика не требуются.")

    doc.save(DOCX_PATH)


def pdf_image(path: Path, width: float) -> PdfImage:
    with Image.open(path) as image:
        ratio = image.height / image.width
    return PdfImage(str(path), width=width, height=width * ratio)


def build_pdf(assets: dict[str, Path]) -> None:
    pdfmetrics.registerFont(TTFont("Arial", "C:/Windows/Fonts/arial.ttf"))
    pdfmetrics.registerFont(TTFont("Arial-Bold", "C:/Windows/Fonts/arialbd.ttf"))
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "BodyRU", parent=styles["BodyText"], fontName="Arial", fontSize=10.5,
        leading=14, textColor=f"#{INK}", spaceAfter=8,
    )
    h1 = ParagraphStyle(
        "H1RU", parent=styles["Heading1"], fontName="Arial-Bold", fontSize=17,
        leading=21, textColor=f"#{ACCENT}", spaceBefore=8, spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "H2RU", parent=styles["Heading2"], fontName="Arial-Bold", fontSize=13,
        leading=16, textColor=f"#{ACCENT}", spaceBefore=8, spaceAfter=6,
    )
    title = ParagraphStyle(
        "TitleRU", parent=styles["Title"], fontName="Arial-Bold", fontSize=23,
        leading=27, textColor=f"#{INK}", alignment=0, spaceAfter=8,
    )
    status = ParagraphStyle(
        "StatusRU", parent=body, fontName="Arial-Bold", textColor=f"#{GREEN}",
        backColor="#E8F5ED", borderPadding=10, spaceBefore=8, spaceAfter=12,
    )
    caption = ParagraphStyle(
        "CaptionRU", parent=body, fontSize=9, textColor=f"#{MUTED}", alignment=1,
    )

    def footer(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont("Arial", 8)
        canvas.setFillColor(f"#{MUTED}")
        canvas.drawString(0.8 * inch, 0.42 * inch, "АП-Риал | Отчёт о выполненных исправлениях")
        canvas.drawRightString(7.7 * inch, 0.42 * inch, str(doc.page))
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(PDF_PATH), pagesize=letter, leftMargin=0.8 * inch, rightMargin=0.8 * inch,
        topMargin=0.72 * inch, bottomMargin=0.65 * inch,
        title="Отчёт об исправлениях на medlic.spb.ru и lfsb.ru",
        author="АП-Риал",
    )
    story = [
        Paragraph("Отчёт об исправлениях на medlic.spb.ru и lfsb.ru", title),
        Paragraph("Повторная проверка и устранение обнаруженных недочётов", h2),
        Paragraph("13 августа 2026 года", body),
        Paragraph("Все перечисленные ниже недочёты устранены и повторно проверены.", status),
        Paragraph(
            "Проведена повторная проверка двух сайтов. Исправлены найденные ошибки в текстах, "
            "мобильной вёрстке и старых подключениях страниц. После публикации сайты заново "
            "проверены на компьютере и телефоне.", body,
        ),
        Paragraph("1. medlic.spb.ru", h1),
        Paragraph(
            "<b>Что было не так:</b> на странице оставались две опечатки, а скрытые пункты "
            "мобильного меню могли увеличивать ширину страницы.", body,
        ),
        Paragraph(
            "<b>Что сделано:</b> «Всеь процессы» заменено на «Все процессы», а "
            "«Росздравнадзоррешает» заменено на «Росздравнадзор решает». Мобильное меню "
            "ограничено шириной экрана. Содержание и внешний вид компьютерной версии не менялись.", body,
        ),
        KeepTogether([pdf_image(assets["medlic_after"], 2.35 * inch), Paragraph("medlic.spb.ru на телефоне после исправления", caption)]),
        PageBreak(),
        Paragraph("2. lfsb.ru", h1),
        Paragraph(
            "<b>Что было не так:</b> сайт был построен на старой фиксированной вёрстке. На "
            "телефоне часть страниц обрезалась, отдельные карточки накладывались друг на друга, "
            "а несколько старых подключений браузер блокировал.", body,
        ),
        Paragraph(
            "<b>Что сделано:</b> все 22 публичные страницы адаптированы для экранов телефона; "
            "шапка, меню, текст, изображения, карточки, боковые блоки и подвал теперь располагаются "
            "последовательно. Старые подключения заменены на рабочие локальные или защищённые версии.", body,
        ),
        Table(
            [[
                [
                    pdf_image(assets["lfsb_before"], 2.05 * inch),
                    Paragraph("До повторной проверки: страница шире экрана", caption),
                ],
                [
                    pdf_image(assets["lfsb_after"], 2.05 * inch),
                    Paragraph("После исправления: содержимое помещается в экран", caption),
                ],
            ]],
            colWidths=[3.25 * inch, 3.25 * inch],
            style=TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ("BOX", (0, 0), (-1, -1), 0, colors.white),
            ]),
        ),
        PageBreak(),
        Paragraph("3. Длинные страницы и карточки", h1),
        Paragraph(
            "Отдельно проверены страницы с большим количеством карточек и старой сложной разметкой. "
            "Карточки выстроены вертикально, текст и изображения больше не перекрывают друг друга.", body,
        ),
        KeepTogether([pdf_image(assets["fstec_after"], 2.15 * inch), Paragraph("Лицензия ФСТЭК", caption)]),
        Spacer(1, 8),
        KeepTogether([pdf_image(assets["kripto_after"], 2.0 * inch), Paragraph("Лицензия на криптографию", caption)]),
        PageBreak(),
        Paragraph("4. Формы и итоговая проверка", h1),
        Paragraph(
            "На обоих сайтах повторно открыты формы «Заказать звонок» и «Задать вопрос» на "
            "компьютере и телефоне. Поля читаемы, имеют одинаковую высоту и не выходят за границы окна.", body,
        ),
        KeepTogether([pdf_image(QA / "lfsb-home-mobile-question.png", 2.4 * inch), Paragraph("Форма на lfsb.ru", caption)]),
        Spacer(1, 8),
        KeepTogether([pdf_image(QA / "medlic-home-mobile-question.png", 2.4 * inch), Paragraph("Форма на medlic.spb.ru", caption)]),
        Paragraph("Что проверено после публикации", h2),
    ]
    for line in (
        "- 22 публичные страницы lfsb.ru на экране телефона 390 px;",
        "- главная lfsb.ru на 320 px и 1440 px;",
        "- главная medlic.spb.ru на 390 px и 1440 px;",
        "- открытие обеих форм на каждом сайте;",
        "- отсутствие обрезки, перекрытий, ошибок браузера и неудачных загрузок.",
    ):
        story.append(Paragraph(line, body))
    story.append(Paragraph("Итог: обнаруженные недочёты устранены. Дополнительные данные от заказчика не требуются.", status))
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def build_audit() -> None:
    results = json.loads((QA / "results.json").read_text(encoding="utf-8"))
    audit = {
        "task_id": "residual-quality-fixes-20260813",
        "created_at": datetime.now(ZoneInfo("Europe/Moscow")).isoformat(timespec="seconds"),
        "status": "passed",
        "client_contact": "not_sent_pending_owner_release",
        "requirements": [
            {"id": "MED-01", "site": "medlic.spb.ru", "status": "passed", "evidence": "qa/medlic-home-desktop.png; qa/medlic-home-mobile.png"},
            {"id": "MED-02", "site": "medlic.spb.ru", "status": "passed", "evidence": "qa/results.json text assertions"},
            {"id": "LFSB-01", "site": "lfsb.ru", "status": "passed", "evidence": "qa/results.json: 22 mobile routes, zero overflow"},
            {"id": "LFSB-02", "site": "lfsb.ru", "status": "passed", "evidence": "qa/lfsb-home-desktop.png; qa/lfsb-home-mobile.png; forms"},
            {"id": "LFSB-03", "site": "lfsb.ru", "status": "passed", "evidence": "qa/lfsb-fstec_dir-mobile.png; qa/lfsb-kripto_dir-mobile.png; qa/lfsb-contakt-mobile.png; qa/lfsb-sendlic-mobile.png"},
        ],
        "qa": {
            "passed": results["passed"],
            "failures": results["failures"],
            "lfsb_checks": len(results["lfsb"]),
            "medlic_checks": len(results["medlic"]),
        },
        "backups": [
            "/home/n/nousroc9/_backups/20260813-202140-residual-quality-fixes",
            "/home/n/nousroc9/_backups/20260813-205104-medlic-mobile-nav-fix",
            "/home/n/nousroc9/_backups/20260813-205419-lfsb-runtime-cleanup",
            "/home/n/nousroc9/_backups/20260813-233000-lfsb-mobile-card-flow-final",
        ],
        "deliverables": {
            "docx": str(DOCX_PATH.relative_to(ROOT)),
            "pdf": str(PDF_PATH.relative_to(ROOT)),
            "audit": str(AUDIT_PATH.relative_to(ROOT)),
        },
    }
    AUDIT_PATH.write_text(json.dumps(audit, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main() -> int:
    assets = prepare_assets()
    build_docx(assets)
    build_pdf(assets)
    build_audit()
    print(json.dumps({"docx": str(DOCX_PATH), "pdf": str(PDF_PATH), "audit": str(AUDIT_PATH)}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
