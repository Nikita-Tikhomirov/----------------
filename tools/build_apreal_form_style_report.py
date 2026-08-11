from __future__ import annotations

import argparse
import hashlib
import json
import textwrap
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DATE = "11 августа 2026 года"
QA_DIR = Path(r"C:\Users\user\AppData\Local\Temp\apreal-form-css-final-v3-20260811")
QA_RESULTS_PATH = QA_DIR / "results.json"
BEFORE_SCREENSHOT = Path(
    r"C:\Users\user\AppData\Local\Temp\apreal-form-css-final-v2-20260811\ed-kgd.ru-desktop-question.png"
)
AFTER_SCREENSHOT = QA_DIR / "ed-kgd.ru-desktop-question.png"
RECIPIENT_MATRIX_PATH = ROOT / "output/ap-real-recipient-matrix-2026-08-11-after-form-css.json"
DEPLOY_MANIFEST_PATH = ROOT / "output/ap-real-form-visual-contract-deploy-2026-08-11.json"

OUTPUT_DOCX = ROOT / "output/documents/AP-Real-client-report-2026-08-11.docx"
OUTPUT_PDF = ROOT / "output/pdf/AP-Real-client-report-2026-08-11.pdf"
OUTPUT_AUDIT = ROOT / "output/ap-real-client-report-audit-2026-08-11.json"
ASSET_DIR = ROOT / "output/ap-real-form-style-report-assets-2026-08-11"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "E8F5EC"
GREEN = "1F7A45"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D8DEE5"
DARK_GRAY = "4B5563"
BLACK = "171717"
WHITE = "FFFFFF"

INTRO_PARAGRAPHS = (
    "В предыдущем результате часть различий между сайтами осталась незамеченной: "
    "проверка опиралась на отдельные примеры, а внешний вид каждой формы на каждом сайте "
    "не был сопоставлен отдельно.",
    "После замечаний проверка проведена заново. На всех 30 сайтах отдельно открыты обе формы "
    "на компьютере и телефоне, исправлены размеры и оформление полей, а настройки адресов "
    "сверены с таблицей клиента.",
)

CORRECTION_ROWS = (
    (
        "Одинаковое оформление полей",
        "Поля имени, телефона, вопроса и проверки должны выглядеть как части одной формы.",
        "На части сайтов высота, отступы, рамки и шрифт у полей различались.",
        "Для всех полей задан единый внешний вид: одинаковая высота, рамка, внутренние отступы и размер текста.",
    ),
    (
        "Нормальная ширина полей",
        "Каждое поле должно занимать доступную ширину формы и оставаться удобным для заполнения.",
        "На сайтах со старой темой подпись ограничивала ширину, поэтому отдельные поля сжимались в узкую колонку.",
        "Ограничение старой темы снято. Подписи и поля занимают всю доступную ширину формы.",
    ),
    (
        "Две формы на каждом сайте",
        "Формы «ЗАКАЗАТЬ ЗВОНОК» и «ЗАДАТЬ ВОПРОС» должны открываться и выглядеть аккуратно.",
        "Ранее одинаковый результат подтверждался не для каждого сайта и не для обеих форм.",
        "Обе формы проверены отдельно на каждом из 30 сайтов. Окна открываются, поля доступны, кнопки и крестик не перекрыты.",
    ),
    (
        "Отображение на телефоне",
        "Форма должна помещаться на экране телефона без обрезки и горизонтальной прокрутки.",
        "Мобильный вид был просмотрен выборочно, поэтому различия между темами сайтов могли остаться незамеченными.",
        "Каждый сайт повторно проверен в мобильном размере. Поля, подписи, капча и кнопка помещаются в окне формы.",
    ),
    (
        "Адреса для заявок",
        "Заявки каждого сайта должны уходить на адрес, указанный клиентом для этого сайта.",
        "Во время диагностики использовался единый контрольный адрес. Оставлять его в рабочей версии было нельзя.",
        "Рабочие адреса возвращены по таблице клиента и проверены в настройках всех форм. Личных и тестовых адресов нет.",
    ),
    (
        "Проверка после публикации",
        "Готовность должна подтверждаться на опубликованных сайтах, а не только по файлам и настройкам.",
        "Предыдущая приемка не обеспечила отдельного зрительного подтверждения для каждой формы и каждого экрана.",
        "После публикации выполнен новый полный просмотр. В приложение включены свежие снимки всех сайтов.",
    ),
)

RESULT_POINTS = (
    "На каждом из 30 сайтов открываются обе формы.",
    "Поля имеют одинаковое оформление и занимают доступную ширину.",
    "На телефоне формы помещаются в окне и не обрезаются.",
    "Кнопки отправки и закрытия доступны, элементы формы не накладываются друг на друга.",
    "В настройках форм сохранены адреса клиента для соответствующих сайтов; личных и тестовых адресов нет.",
)

METHOD_POINTS = (
    "Каждый сайт открыт в свежей сессии после публикации.",
    "На каждом сайте отдельно открыты формы обратного звонка и вопроса.",
    "Проверка повторена в размере экрана компьютера и телефона.",
    "Проверены размеры полей, подписи, капча, кнопка отправки, крестик и отсутствие наложений.",
    "Адреса получателей сверены непосредственно в серверных настройках форм.",
)

CLIENT_NOTE = (
    "В ходе этой повторной проверки тестовые заявки в рабочие ящики не отправлялись, чтобы не создавать "
    "лишние обращения. Проверены внешний вид и работа элементов формы в браузере, а также адреса получателей "
    "в серверных настройках."
)


def collect_client_visible_text() -> str:
    chunks: list[str] = [
        "Что требовалось",
        "Что было сделано неправильно",
        "Что исправлено сейчас",
        *INTRO_PARAGRAPHS,
        *RESULT_POINTS,
        *METHOD_POINTS,
        CLIENT_NOTE,
    ]
    for row in CORRECTION_ROWS:
        chunks.extend(row)
    return " ".join(chunks)


def build_evidence_batches(domains: list[str], batch_size: int = 6) -> list[list[str]]:
    if batch_size < 1:
        raise ValueError("batch_size must be positive")
    return [domains[index : index + batch_size] for index in range(0, len(domains), batch_size)]


def site_page_facts(domain: str, recipient: str) -> tuple[str, ...]:
    return (
        f"Сайт: {domain}",
        "Обе формы открываются: «ЗАКАЗАТЬ ЗВОНОК» и «ЗАДАТЬ ВОПРОС».",
        "Поля приведены к единому виду и занимают доступную ширину формы.",
        "Обе формы проверены на компьютере и телефоне после публикации.",
        f"Адрес для заявок в рабочих настройках: {recipient}.",
        "Личных и тестовых адресов нет.",
    )


def load_qa_results(path: Path) -> dict[tuple[str, str], dict[str, Any]]:
    raw = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(raw, list) or not raw:
        raise ValueError("QA results must be a non-empty list")

    index: dict[tuple[str, str], dict[str, Any]] = {}
    errors: list[str] = []
    for row in raw:
        key = (str(row.get("domain", "")).strip(), str(row.get("viewport", "")).strip())
        if not all(key):
            raise ValueError("QA result is missing domain or viewport")
        if key in index:
            raise ValueError(f"Duplicate QA result: {key[0]} / {key[1]}")
        index[key] = row

        row_errors: list[str] = []
        if int(row.get("status") or 0) != 200:
            row_errors.append(f"HTTP {row.get('status')}")
        for field in ("failures", "pageErrors", "criticalConsoleErrors", "requestFailures"):
            if row.get(field):
                row_errors.append(f"{field}: {row[field]}")
        forms = row.get("forms") or {}
        actions = row.get("actions") or {}
        captured_kinds = set(forms) | set(actions)
        if not all(kind in captured_kinds for kind in ("callback", "question")):
            row_errors.append("both forms were not captured")
        if row_errors:
            errors.append(f"{key[0]} / {key[1]}: {'; '.join(row_errors)}")

    if errors:
        raise ValueError("QA failures: " + " | ".join(errors))

    domains = {domain for domain, _ in index}
    for domain in domains:
        viewports = {viewport for candidate, viewport in index if candidate == domain}
        if viewports != {"desktop", "mobile"}:
            raise ValueError(f"Each site must have desktop and mobile QA: {domain}")
    return index


def load_recipient_matrix(path: Path) -> dict[str, Any]:
    matrix = json.loads(path.read_text(encoding="utf-8"))
    summary = matrix.get("summary") or {}
    checks = matrix.get("checks") or []
    if not summary.get("complete"):
        raise ValueError("Recipient matrix is incomplete")
    if summary.get("failed"):
        raise ValueError(f"Recipient matrix has failed checks: {summary['failed']}")
    if summary.get("personal_recipient_hits"):
        raise ValueError("Recipient matrix contains personal or test recipients")
    if int(summary.get("checks") or 0) != int(summary.get("passed") or -1):
        raise ValueError("Recipient matrix pass count does not match check count")
    if not checks or any(not check.get("passed") for check in checks):
        raise ValueError("Recipient matrix includes an unconfirmed route")
    return matrix


def report_site_order(
    qa_index: dict[tuple[str, str], dict[str, Any]], matrix: dict[str, Any]
) -> list[str]:
    qa_domains = {domain for domain, _ in qa_index}
    route_domains = set((matrix.get("expected_sites") or {}).keys())
    if qa_domains != route_domains:
        missing_qa = sorted(route_domains - qa_domains)
        missing_routes = sorted(qa_domains - route_domains)
        raise ValueError(
            "QA and recipient coverage differs: "
            f"missing QA={missing_qa}; missing routes={missing_routes}"
        )
    return sorted(qa_domains)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def _draw_wrapped(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str,
    max_width: int,
    line_spacing: int = 8,
) -> int:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if current and width > max_width:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)

    x, y = xy
    for line in lines:
        bbox = draw.textbbox((x, y), line, font=font)
        draw.text((x, y), line, font=font, fill=fill)
        y += bbox[3] - bbox[1] + line_spacing
    return y


def _paste_contained(canvas: Image.Image, source_path: Path, box: tuple[int, int, int, int]) -> None:
    if not source_path.exists():
        raise FileNotFoundError(source_path)
    x1, y1, x2, y2 = box
    with Image.open(source_path) as opened:
        image = opened.convert("RGB")
    image.thumbnail((x2 - x1, y2 - y1), Image.Resampling.LANCZOS)
    x = x1 + ((x2 - x1) - image.width) // 2
    y = y1 + ((y2 - y1) - image.height) // 2
    canvas.paste(image, (x, y))


def _screenshot_path(domain: str, viewport: str, kind: str) -> Path:
    return QA_DIR / f"{domain}-{viewport}-{kind}.png"


def validate_screenshot_evidence(domains: Iterable[str]) -> list[Path]:
    paths: list[Path] = [BEFORE_SCREENSHOT, AFTER_SCREENSHOT]
    for domain in domains:
        for viewport in ("desktop", "mobile"):
            for kind in ("callback", "question"):
                paths.append(_screenshot_path(domain, viewport, kind))
    missing = [str(path) for path in paths if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing report screenshots: " + ", ".join(missing))
    return paths


def build_before_after_asset() -> Path:
    output = ASSET_DIR / "before-after-ed-kgd.png"
    canvas = Image.new("RGB", (2400, 1660), f"#{WHITE}")
    draw = ImageDraw.Draw(canvas)
    draw.text((90, 55), "Пример исправления: форма вопроса на ed-kgd.ru", font=_font(52, True), fill=f"#{DARK_BLUE}")
    labels = (
        ("До исправления: поля сжаты старой темой", BEFORE_SCREENSHOT),
        ("После исправления: поля занимают ширину формы", AFTER_SCREENSHOT),
    )
    for index, (label, source) in enumerate(labels):
        x1 = 90 + index * 1150
        x2 = x1 + 1060
        draw.rounded_rectangle((x1, 150, x2, 1570), radius=18, outline=f"#{MID_GRAY}", width=4, fill="#FAFBFC")
        _draw_wrapped(draw, (x1 + 32, 185), label, _font(34, True), f"#{BLACK}", 990, 7)
        _paste_contained(canvas, source, (x1 + 28, 270, x2 - 28, 1530))
    output.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(output, quality=94)
    return output


def build_representative_asset(domains: list[str], viewport: str) -> Path:
    output = ASSET_DIR / f"representative-{viewport}.png"
    canvas = Image.new("RGB", (2400, 3000), f"#{WHITE}")
    draw = ImageDraw.Draw(canvas)
    title = "Формы на компьютере" if viewport == "desktop" else "Формы на телефоне"
    draw.text((90, 55), title, font=_font(54, True), fill=f"#{DARK_BLUE}")
    draw.text(
        (90, 130),
        "Слева — обратный звонок, справа — форма вопроса. Изображения показаны целиком.",
        font=_font(29),
        fill=f"#{DARK_GRAY}",
    )
    row_height = 675
    for index, domain in enumerate(domains):
        top = 215 + index * row_height
        draw.rounded_rectangle((80, top, 2320, top + 620), radius=16, outline=f"#{MID_GRAY}", width=3)
        draw.text((115, top + 22), domain, font=_font(34, True), fill=f"#{BLACK}")
        _paste_contained(canvas, _screenshot_path(domain, viewport, "callback"), (115, top + 82, 1170, top + 590))
        _paste_contained(canvas, _screenshot_path(domain, viewport, "question"), (1230, top + 82, 2285, top + 590))
    output.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(output, quality=94)
    return output


def build_full_evidence_asset(domains: list[str], page_number: int) -> Path:
    output = ASSET_DIR / f"all-sites-{page_number:02d}.png"
    canvas = Image.new("RGB", (2480, 3508), f"#{WHITE}")
    draw = ImageDraw.Draw(canvas)
    draw.text(
        (100, 60),
        f"Свежая проверка всех сайтов — лист {page_number}",
        font=_font(52, True),
        fill=f"#{DARK_BLUE}",
    )
    draw.text(
        (100, 135),
        "Для каждого сайта: компьютер и телефон, обратный звонок и форма вопроса.",
        font=_font(28),
        fill=f"#{DARK_GRAY}",
    )

    row_height = 530
    for row_index, domain in enumerate(domains):
        top = 215 + row_index * row_height
        draw.rounded_rectangle((85, top, 2395, top + 485), radius=14, outline=f"#{MID_GRAY}", width=3)
        draw.text((110, top + 16), domain, font=_font(31, True), fill=f"#{BLACK}")
        labels = (
            ("Компьютер · звонок", "desktop", "callback"),
            ("Компьютер · вопрос", "desktop", "question"),
            ("Телефон · звонок", "mobile", "callback"),
            ("Телефон · вопрос", "mobile", "question"),
        )
        for column, (label, viewport, kind) in enumerate(labels):
            left = 110 + column * 570
            draw.text((left, top + 63), label, font=_font(22, True), fill=f"#{DARK_GRAY}")
            _paste_contained(
                canvas,
                _screenshot_path(domain, viewport, kind),
                (left, top + 100, left + 535, top + 455),
            )
    output.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(output, quality=92)
    return output


def build_site_evidence_asset(domain: str, recipient: str, page_number: int) -> Path:
    safe_domain = domain.replace(".", "-")
    output = ASSET_DIR / f"site-{page_number:02d}-{safe_domain}.jpg"
    canvas = Image.new("RGB", (2480, 3508), f"#{WHITE}")
    draw = ImageDraw.Draw(canvas)
    draw.text((100, 60), domain, font=_font(58, True), fill=f"#{DARK_BLUE}")
    draw.text(
        (100, 140),
        f"Проверка после исправления · сайт {page_number} из 30",
        font=_font(29),
        fill=f"#{DARK_GRAY}",
    )

    facts = site_page_facts(domain, recipient)[1:]
    fact_y = 215
    for fact in facts:
        draw.ellipse((105, fact_y + 9, 121, fact_y + 25), fill=f"#{BLUE}")
        fact_y = _draw_wrapped(
            draw,
            (145, fact_y),
            fact,
            _font(27),
            f"#{BLACK}",
            2200,
            7,
        ) + 8

    cards = (
        ("Компьютер · обратный звонок", "desktop", "callback", 90, 570),
        ("Компьютер · форма вопроса", "desktop", "question", 1270, 570),
        ("Телефон · обратный звонок", "mobile", "callback", 90, 2020),
        ("Телефон · форма вопроса", "mobile", "question", 1270, 2020),
    )
    for label, viewport, kind, left, top in cards:
        right = left + 1120
        bottom = top + 1370
        draw.rounded_rectangle(
            (left, top, right, bottom),
            radius=18,
            outline=f"#{MID_GRAY}",
            width=4,
            fill="#FAFBFC",
        )
        draw.text((left + 28, top + 24), label, font=_font(29, True), fill=f"#{BLACK}")
        _paste_contained(
            canvas,
            _screenshot_path(domain, viewport, kind),
            (left + 28, top + 90, right - 28, bottom - 28),
        )

    output.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(output, format="JPEG", quality=84, optimize=True, subsampling=1)
    return output


def build_assets(domains: list[str], matrix: dict[str, Any]) -> dict[str, Any]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    representative = ["apreal.ru", "mca24.ru", "nousro-spb.ru", "license39.ru"]
    missing_representative = sorted(set(representative) - set(domains))
    if missing_representative:
        raise ValueError(f"Representative sites are missing: {missing_representative}")
    evidence_batches = build_evidence_batches(domains, batch_size=6)
    expected_sites = matrix["expected_sites"]
    return {
        "before_after": build_before_after_asset(),
        "desktop": build_representative_asset(representative, "desktop"),
        "mobile": build_representative_asset(representative, "mobile"),
        "per_site": {
            domain: build_site_evidence_asset(domain, expected_sites[domain], page_number=index + 1)
            for index, domain in enumerate(domains)
        },
        "evidence_batches": evidence_batches,
    }


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _add_page_number(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def _configure_document(document: Document) -> None:
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run("ГК «АП-Риал»  ·  ")
        footer_run.font.name = "Calibri"
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
        _add_page_number(footer)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 24, DARK_BLUE),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, DARK_BLUE),
        ("Heading 3", 11.5, DARK_BLUE),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)


def _docx_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)


def _docx_status_box(document: Document) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    _set_cell_shading(cell, PALE_GREEN)
    _set_cell_margins(cell, 170, 180, 170, 180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Итог: последние замечания по формам устранены и повторно проверены на всех 30 сайтах.")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)


def _docx_correction_card(document: Document, row: tuple[str, str, str, str]) -> None:
    title, required, wrong, fixed = row
    table = document.add_table(rows=4, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(5.1)
    labels = (
        (title, "", PALE_BLUE),
        ("Что требовалось", required, LIGHT_GRAY),
        ("Что было сделано неправильно", wrong, "FFF4DE"),
        ("Что исправлено сейчас", fixed, PALE_GREEN),
    )
    for index, (label, value, fill) in enumerate(labels):
        left = table.cell(index, 0)
        right = table.cell(index, 1)
        _set_cell_shading(left, fill)
        _set_cell_shading(right, fill)
        _set_cell_margins(left)
        _set_cell_margins(right)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if index == 0:
            merged = left.merge(right)
            paragraph = merged.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(label)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        else:
            left.paragraphs[0].add_run(label).bold = True
            right.paragraphs[0].add_run(value)
            left.paragraphs[0].paragraph_format.space_after = Pt(0)
            right.paragraphs[0].paragraph_format.space_after = Pt(0)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def build_docx(domains: list[str], matrix: dict[str, Any], assets: dict[str, Any]) -> Path:
    document = Document()
    _configure_document(document)
    properties = document.core_properties
    properties.title = "Исправление форм на сайтах ГК «АП-Риал»"
    properties.subject = "Отчет по устранению последних замечаний"
    properties.author = "Никита Тихонов"
    properties.keywords = "АП-Риал, формы, проверка сайтов"

    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = eyebrow.add_run("ОТЧЕТ ПО УСТРАНЕНИЮ ПОСЛЕДНИХ ЗАМЕЧАНИЙ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    title = document.add_paragraph(style="Title")
    title.add_run("Исправление форм на сайтах\nГК «АП-Риал»")
    date = document.add_paragraph(REPORT_DATE)
    date.runs[0].font.color.rgb = RGBColor.from_string(DARK_GRAY)
    date.paragraph_format.space_after = Pt(18)

    _docx_status_box(document)
    document.add_paragraph()
    for paragraph_text in INTRO_PARAGRAPHS:
        document.add_paragraph(paragraph_text)

    document.add_heading("Что подтверждено сейчас", level=1)
    for point in RESULT_POINTS:
        _docx_bullet(document, point)

    note_table = document.add_table(rows=1, cols=1)
    note_cell = note_table.cell(0, 0)
    _set_cell_shading(note_cell, PALE_BLUE)
    _set_cell_margins(note_cell, 150, 170, 150, 170)
    note_cell.paragraphs[0].add_run(CLIENT_NOTE)

    document.add_page_break()
    document.add_heading("Что было исправлено", level=1)
    document.add_paragraph(
        "Ниже по каждому замечанию указано требование, ошибка предыдущего результата и фактическое исправление."
    )
    for index, row in enumerate(CORRECTION_ROWS):
        _docx_correction_card(document, row)
        if index == 2:
            document.add_page_break()
            document.add_heading("Что было исправлено — продолжение", level=1)

    document.add_page_break()
    document.add_heading("Наглядный пример исправления", level=1)
    document.add_paragraph(
        "На снимке слева видно прежнее сжатие полей старой темой. Справа — опубликованный вариант после исправления."
    )
    document.add_picture(str(assets["before_after"]), width=Inches(6.75))

    for key, heading, caption in (
        (
            "desktop",
            "Примеры после исправления — компьютер",
            "Показаны разные группы сайтов. Обе формы открыты на опубликованных страницах после исправления.",
        ),
        (
            "mobile",
            "Примеры после исправления — телефон",
            "Формы помещаются в мобильном окне, поля не сжаты и не обрезаны.",
        ),
    ):
        document.add_page_break()
        document.add_heading(heading, level=1)
        document.add_paragraph(caption)
        document.add_picture(str(assets[key]), width=Inches(6.15))

    document.add_page_break()
    document.add_heading("Как проведена повторная проверка", level=1)
    for point in METHOD_POINTS:
        _docx_bullet(document, point)
    document.add_heading("Сайты и адреса для заявок", level=1)
    document.add_paragraph(
        "Ниже перечислены все сайты из этой проверки и адреса, сохраненные в их рабочих настройках."
    )
    route_table = document.add_table(rows=1, cols=2)
    route_table.autofit = False
    route_table.columns[0].width = Inches(2.45)
    route_table.columns[1].width = Inches(4.2)
    headers = route_table.rows[0].cells
    headers[0].text = "Сайт"
    headers[1].text = "Адрес для заявок"
    for cell in headers:
        _set_cell_shading(cell, DARK_BLUE)
        _set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    expected_sites = matrix["expected_sites"]
    for domain in domains:
        cells = route_table.add_row().cells
        cells[0].text = domain
        cells[1].text = expected_sites[domain]
        for cell in cells:
            _set_cell_margins(cell, 70, 100, 70, 100)
        if len(route_table.rows) % 2 == 0:
            _set_cell_shading(cells[0], LIGHT_GRAY)
            _set_cell_shading(cells[1], LIGHT_GRAY)

    for domain in domains:
        document.add_page_break()
        evidence_path = assets["per_site"][domain]
        document.add_picture(str(evidence_path), width=Inches(6.55))

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


def _pdf_font_paths() -> tuple[Path, Path]:
    regular = Path(r"C:\Windows\Fonts\arial.ttf")
    bold = Path(r"C:\Windows\Fonts\arialbd.ttf")
    if not regular.exists() or not bold.exists():
        raise FileNotFoundError("Arial fonts are required for the PDF report")
    return regular, bold


def build_pdf(domains: list[str], matrix: dict[str, Any], assets: dict[str, Any]) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        Image as PdfImage,
        KeepTogether,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    regular, bold = _pdf_font_paths()
    pdfmetrics.registerFont(TTFont("ReportArial", str(regular)))
    pdfmetrics.registerFont(TTFont("ReportArialBold", str(bold)))

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "ReportBody",
        parent=styles["BodyText"],
        fontName="ReportArial",
        fontSize=9.4,
        leading=12.2,
        textColor=colors.HexColor(f"#{BLACK}"),
        spaceAfter=7,
    )
    small = ParagraphStyle(
        "ReportSmall",
        parent=body,
        fontSize=7.8,
        leading=9.8,
        spaceAfter=0,
    )
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=body,
        fontName="ReportArialBold",
        fontSize=23,
        leading=27,
        textColor=colors.HexColor(f"#{DARK_BLUE}"),
        spaceAfter=12,
    )
    eyebrow = ParagraphStyle(
        "ReportEyebrow",
        parent=body,
        fontName="ReportArialBold",
        fontSize=8,
        leading=10,
        textColor=colors.HexColor(f"#{BLUE}"),
        spaceAfter=8,
    )
    heading = ParagraphStyle(
        "ReportHeading",
        parent=body,
        fontName="ReportArialBold",
        fontSize=15,
        leading=18,
        textColor=colors.HexColor(f"#{BLUE}"),
        spaceBefore=4,
        spaceAfter=8,
    )
    heading2 = ParagraphStyle(
        "ReportHeading2",
        parent=body,
        fontName="ReportArialBold",
        fontSize=11,
        leading=13,
        textColor=colors.HexColor(f"#{DARK_BLUE}"),
        spaceAfter=5,
    )
    label_style = ParagraphStyle(
        "ReportLabel",
        parent=small,
        fontName="ReportArialBold",
        textColor=colors.HexColor(f"#{DARK_BLUE}"),
    )
    header_style = ParagraphStyle(
        "ReportHeader",
        parent=small,
        fontName="ReportArialBold",
        textColor=colors.white,
    )

    def footer(canvas: Any, document: Any) -> None:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor(f"#{MID_GRAY}"))
        canvas.line(0.65 * inch, 0.48 * inch, 7.85 * inch, 0.48 * inch)
        canvas.setFont("ReportArial", 7.5)
        canvas.setFillColor(colors.HexColor(f"#{DARK_GRAY}"))
        canvas.drawString(0.65 * inch, 0.3 * inch, "ГК «АП-Риал»")
        canvas.drawRightString(7.85 * inch, 0.3 * inch, str(document.page))
        canvas.restoreState()

    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=letter,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.62 * inch,
        bottomMargin=0.62 * inch,
        title="Исправление форм на сайтах ГК «АП-Риал»",
        author="Никита Тихонов",
    )
    story: list[Any] = []
    story.append(Paragraph("ОТЧЕТ ПО УСТРАНЕНИЮ ПОСЛЕДНИХ ЗАМЕЧАНИЙ", eyebrow))
    story.append(Paragraph("Исправление форм на сайтах<br/>ГК «АП-Риал»", title_style))
    story.append(Paragraph(REPORT_DATE, body))
    story.append(Spacer(1, 8))
    status = Table(
        [[Paragraph("<b>Итог:</b> последние замечания по формам устранены и повторно проверены на всех 30 сайтах.", body)]],
        colWidths=[7.2 * inch],
    )
    status.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{PALE_GREEN}")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(f"#{GREEN}")),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
            ]
        )
    )
    story.extend([status, Spacer(1, 13)])
    for paragraph_text in INTRO_PARAGRAPHS:
        story.append(Paragraph(paragraph_text, body))
    story.append(Paragraph("Что подтверждено сейчас", heading))
    for point in RESULT_POINTS:
        story.append(Paragraph(f"• {point}", body))
    note = Table([[Paragraph(CLIENT_NOTE, small)]], colWidths=[7.2 * inch])
    note.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{PALE_BLUE}")),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(f"#{BLUE}")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.extend([Spacer(1, 8), note, PageBreak()])

    story.append(Paragraph("Что было исправлено", heading))
    story.append(
        Paragraph(
            "По каждому замечанию ниже указано требование, ошибка предыдущего результата и фактическое исправление.",
            body,
        )
    )
    for index, (card_title, required, wrong, fixed) in enumerate(CORRECTION_ROWS):
        card_data = [
            [Paragraph(card_title, heading2), ""],
            [Paragraph("Что требовалось", label_style), Paragraph(required, small)],
            [Paragraph("Что было сделано неправильно", label_style), Paragraph(wrong, small)],
            [Paragraph("Что исправлено сейчас", label_style), Paragraph(fixed, small)],
        ]
        card = Table(card_data, colWidths=[1.75 * inch, 5.45 * inch], hAlign=TA_LEFT)
        card.setStyle(
            TableStyle(
                [
                    ("SPAN", (0, 0), (1, 0)),
                    ("BACKGROUND", (0, 0), (1, 0), colors.HexColor(f"#{PALE_BLUE}")),
                    ("BACKGROUND", (0, 1), (1, 1), colors.HexColor(f"#{LIGHT_GRAY}")),
                    ("BACKGROUND", (0, 2), (1, 2), colors.HexColor("#FFF4DE")),
                    ("BACKGROUND", (0, 3), (1, 3), colors.HexColor(f"#{PALE_GREEN}")),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{MID_GRAY}")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 7),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.append(KeepTogether([card, Spacer(1, 8)]))
        if index == 2:
            story.extend([PageBreak(), Paragraph("Что было исправлено — продолжение", heading)])

    story.extend([PageBreak(), Paragraph("Наглядный пример исправления", heading)])
    story.append(
        Paragraph(
            "Слева видно прежнее сжатие полей старой темой. Справа — опубликованный вариант после исправления.",
            body,
        )
    )
    story.append(PdfImage(str(assets["before_after"]), width=7.2 * inch, height=4.98 * inch))

    for key, heading_text, caption in (
        (
            "desktop",
            "Примеры после исправления — компьютер",
            "Показаны разные группы сайтов. Обе формы открыты на опубликованных страницах после исправления.",
        ),
        (
            "mobile",
            "Примеры после исправления — телефон",
            "Формы помещаются в мобильном окне, поля не сжаты и не обрезаны.",
        ),
    ):
        story.extend([PageBreak(), Paragraph(heading_text, heading), Paragraph(caption, body)])
        story.append(PdfImage(str(assets[key]), width=6.0 * inch, height=7.5 * inch))

    story.extend([PageBreak(), Paragraph("Как проведена повторная проверка", heading)])
    for point in METHOD_POINTS:
        story.append(Paragraph(f"• {point}", body))
    story.append(Paragraph("Сайты и адреса для заявок", heading))
    story.append(
        Paragraph(
            "Ниже перечислены все сайты из этой проверки и адреса, сохраненные в их рабочих настройках.",
            body,
        )
    )
    route_data = [[Paragraph("Сайт", header_style), Paragraph("Адрес для заявок", header_style)]]
    expected_sites = matrix["expected_sites"]
    route_data.extend(
        [Paragraph(domain, small), Paragraph(expected_sites[domain], small)] for domain in domains
    )
    route_table = Table(route_data, colWidths=[2.35 * inch, 4.85 * inch], repeatRows=1)
    route_style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{DARK_BLUE}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{MID_GRAY}")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 3.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5),
    ]
    for row_index in range(2, len(route_data), 2):
        route_style.append(("BACKGROUND", (0, row_index), (-1, row_index), colors.HexColor(f"#{LIGHT_GRAY}")))
    route_table.setStyle(TableStyle(route_style))
    story.append(route_table)

    for domain in domains:
        story.append(PageBreak())
        evidence_path = assets["per_site"][domain]
        story.append(PdfImage(str(evidence_path), width=6.58 * inch, height=9.31 * inch))

    document.build(story, onFirstPage=footer, onLaterPages=footer)
    return OUTPUT_PDF


def _load_deploy_manifest(path: Path) -> dict[str, Any]:
    manifest = json.loads(path.read_text(encoding="utf-8"))
    if manifest.get("rollback_performed"):
        raise ValueError("The live deployment was rolled back")
    published = manifest.get("published") or manifest.get("sites") or []
    if not published:
        raise ValueError("Deployment manifest has no published sites")
    return manifest


def _pdf_page_count(path: Path) -> int:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is required to finalize the report audit") from exc
    return len(PdfReader(str(path)).pages)


def write_audit(
    domains: list[str],
    matrix: dict[str, Any],
    assets: dict[str, Any],
    deploy_manifest: dict[str, Any],
    visual_review_manifest: Path | None = None,
) -> dict[str, Any]:
    page_count = _pdf_page_count(OUTPUT_PDF)
    visual_review: dict[str, Any] = {
        "result": "pending",
        "all_pages_reviewed": False,
        "manifest": None,
    }
    if visual_review_manifest is not None:
        review = json.loads(visual_review_manifest.read_text(encoding="utf-8"))
        reviewed_pages = sorted(set(int(page) for page in review.get("client_report_pages", [])))
        if reviewed_pages != list(range(1, page_count + 1)):
            raise ValueError(
                f"Visual review does not cover every PDF page: expected 1..{page_count}, got {reviewed_pages}"
            )
        visual_review = {
            "result": "passed",
            "all_pages_reviewed": True,
            "manifest": str(visual_review_manifest),
            "reviewed_at": review.get("reviewed_at"),
            "reviewer": review.get("reviewer"),
        }

    audit = {
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "scope": "latest_form_corrections",
        "sites": domains,
        "site_count": len(domains),
        "qa": {
            "source": str(QA_RESULTS_PATH),
            "views": len(domains) * 2,
            "desktop_mobile_per_site": True,
            "callback_and_question_per_view": True,
            "failed_views": [],
        },
        "recipient_settings": {
            "source": str(RECIPIENT_MATRIX_PATH),
            "checks": matrix["summary"]["checks"],
            "failed": matrix["summary"]["failed"],
            "personal_recipient_hits": matrix["summary"]["personal_recipient_hits"],
            "mailbox_delivery_claimed": False,
        },
        "deployment": {
            "source": str(DEPLOY_MANIFEST_PATH),
            "backup": (
                deploy_manifest.get("backup")
                or deploy_manifest.get("backup_dir")
                or deploy_manifest.get("backup_root")
            ),
            "rollback_performed": False,
        },
        "report": {
            "docx": str(OUTPUT_DOCX),
            "docx_sha256": _sha256(OUTPUT_DOCX),
            "pdf": str(OUTPUT_PDF),
            "pdf_sha256": _sha256(OUTPUT_PDF),
            "pdf_pages": page_count,
            "asset_dir": str(ASSET_DIR),
            "evidence_batches": assets["evidence_batches"],
            "site_evidence": {domain: str(assets["per_site"][domain]) for domain in domains},
        },
        "visual_review": visual_review,
        "client_contact": {
            "performed": False,
            "release_required": True,
        },
    }
    OUTPUT_AUDIT.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    return audit


def finalize_existing_visual_review(manifest_path: Path) -> dict[str, Any]:
    if not OUTPUT_AUDIT.exists() or not OUTPUT_PDF.exists():
        raise FileNotFoundError("The generated report and audit must exist before visual review finalization")
    audit = json.loads(OUTPUT_AUDIT.read_text(encoding="utf-8"))
    review = json.loads(manifest_path.read_text(encoding="utf-8"))
    current_hash = _sha256(OUTPUT_PDF)
    reviewed_hash = str(review.get("source_pdf_sha256", "")).casefold()
    if reviewed_hash != current_hash.casefold():
        raise ValueError("Visual review PDF hash does not match the current report")

    page_count = _pdf_page_count(OUTPUT_PDF)
    reviewed_pages = sorted(set(int(page) for page in review.get("client_report_pages", [])))
    if reviewed_pages != list(range(1, page_count + 1)):
        raise ValueError(
            f"Visual review does not cover every PDF page: expected 1..{page_count}, got {reviewed_pages}"
        )

    audit["visual_review"] = {
        "result": "passed",
        "all_pages_reviewed": True,
        "manifest": str(manifest_path),
        "reviewed_at": review.get("reviewed_at"),
        "reviewer": review.get("reviewer"),
        "source_pdf_sha256": current_hash,
        "checks": review.get("checks", {}),
    }
    OUTPUT_AUDIT.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    return audit


def build_report(visual_review_manifest: Path | None = None) -> dict[str, Any]:
    qa_index = load_qa_results(QA_RESULTS_PATH)
    matrix = load_recipient_matrix(RECIPIENT_MATRIX_PATH)
    domains = report_site_order(qa_index, matrix)
    if len(domains) != 30:
        raise ValueError(f"Expected 30 report sites, got {len(domains)}")
    validate_screenshot_evidence(domains)
    deploy_manifest = _load_deploy_manifest(DEPLOY_MANIFEST_PATH)
    assets = build_assets(domains, matrix)
    build_docx(domains, matrix, assets)
    build_pdf(domains, matrix, assets)
    audit = write_audit(domains, matrix, assets, deploy_manifest, visual_review_manifest)
    return {
        "docx": str(OUTPUT_DOCX),
        "pdf": str(OUTPUT_PDF),
        "audit": str(OUTPUT_AUDIT),
        "pages": audit["report"]["pdf_pages"],
        "visual_review": audit["visual_review"],
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the AP-Real form correction client report")
    parser.add_argument(
        "--visual-review-manifest",
        type=Path,
        help="JSON manifest listing every visually reviewed PDF page",
    )
    parser.add_argument(
        "--finalize-only",
        action="store_true",
        help="Finalize the existing audit without rebuilding the reviewed report",
    )
    args = parser.parse_args()
    if args.finalize_only:
        if args.visual_review_manifest is None:
            parser.error("--finalize-only requires --visual-review-manifest")
        audit = finalize_existing_visual_review(args.visual_review_manifest)
        print(json.dumps({"audit": str(OUTPUT_AUDIT), "visual_review": audit["visual_review"]}, ensure_ascii=False, indent=2))
        return 0
    result = build_report(args.visual_review_manifest)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
