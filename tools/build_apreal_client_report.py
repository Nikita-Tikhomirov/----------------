from __future__ import annotations

import argparse
import hashlib
import json
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from tools import build_apreal_acceptance_report as evidence
except ModuleNotFoundError:
    import build_apreal_acceptance_report as evidence  # type: ignore[no-redef]


ROOT = Path(__file__).resolve().parents[1]
REPORT_DATE = "05.08.2026"
OUTPUT_DOCX = ROOT / "output/documents/AP-Real-client-report-2026-08-05.docx"
OUTPUT_PDF = ROOT / "output/pdf/AP-Real-client-report-2026-08-05.pdf"
COVER_NOTE_DOCX = ROOT / "output/documents/AP-Real-cover-note-draft-2026-08-05.docx"
COVER_NOTE_PDF = ROOT / "output/pdf/AP-Real-cover-note-draft-2026-08-05.pdf"
INTERNAL_NOTE = ROOT / "output/AP-Real-internal-incident-note-2026-08-05.md"
OUTPUT_AUDIT = ROOT / "output/ap-real-client-report-audit-2026-08-05.json"
ASSET_DIR = ROOT / "output/ap-real-client-report-assets-2026-08-05"

FRESH_QA_DIR = ROOT / "output/ap-real-owner-report-recheck-2026-08-05"
FRESH_QA_RESULTS = FRESH_QA_DIR / "results.json"
POST_CORRECTION_QA_DIR = ROOT / "output/ap-real-main-chrome-responsive-final-native-click-2026-08-03"
FORM_BOARD_DIR = ROOT / "output/ap-real-final-client-visual-review-2026-08-02"
MIGRATION_QA_DIR = ROOT / "output/ap-real-migration-qa-2026-08-02"
MIGRATION_QA_RESULTS = MIGRATION_QA_DIR / "final-results.json"
MCHS_VRN_RESULTS = MIGRATION_QA_DIR / "mchs-vrn-staged-results.json"
VIDEO_QA_DIR = ROOT / "output/ap-real-video-review-2026-08-02"
MAIL_EVIDENCE_DIR = ROOT / "output/ap-real-evidence-2026-08-02"
SENDER_DELIVERY_PATH = ROOT / "output/ap-real-post-send-form-submissions-2026-08-02.json"
RECIPIENT_MATRIX_PATH = ROOT / "output/ap-real-recipient-matrix-2026-08-05.json"
MAILBOX_RECEIPT_PATH = ROOT / "output/ap-real-post-send-main-mailru-accounts-2026-08-02.json"
MAILBOX_RECEIPT_SCREENSHOT = ROOT / "output/ap-real-post-send-main-mailru-accounts-2026-08-02.png"
HIDDEN_VIDEO_EVIDENCE_PATH = ROOT / "output/ap-real-hidden-video-live-check-2026-08-05.json"
MAIL_DELIVERY_SCOPE = "mailbox_confirmed_sites_only"
MAILBOX_CONFIRMED_SITES = ("medlic.spb.ru",)
MAILBOX_RECEIPT_MARKERS = {
    "APREAL-POST-SEND-20260802-1900-medlic.spb.ru-callback",
    "APREAL-POST-SEND-20260802-1900-medlic.spb.ru-question",
}

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E0E7"
DARK_GRAY = "4B5563"
GREEN = "1F7A45"
PALE_GREEN = "E8F5EC"
AMBER = "A15C00"
PALE_AMBER = "FFF4DE"
RED = "9B1C1C"
PALE_RED = "FDECEC"
BLACK = "111111"

MIGRATIONS_LIVE = [
    "91web.ru",
    "apreal-nn.ru",
    "apreal-volgograd.ru",
    "apreal.ru",
    "apreal36.ru",
    "apreal72.ru",
    "docp.ru",
    "elecktro.ru",
    "fsa-lab.ru",
    "fste.ru",
    "lfsb.ru",
    "mca24.ru",
    "med-license.ru",
    "mhsl.ru",
    "moopb.ru",
    "otxodi.ru",
    "rectavr.ru",
    "shopap.ru",
    "apreal.spb.ru",
    "fstek.spb.ru",
    "medlic.spb.ru",
    "electro-reg.ru",
    "license39.ru",
    "mchs78.ru",
    "medtex39.ru",
    "medtex78.ru",
    "minkult78.ru",
    "mchs-spb.ru",
]
MIGRATIONS_SOURCE_PLACEHOLDER = ["othodi-spb.ru"]
MIGRATIONS_STAGED = ["mchs-vrn.ru"]
MIGRATIONS_BLOCKED = [
    "dpocenter.ru",
    "feo-edem.ru",
    "linkedin.com.moopb.ru",
    "aklab-spb.ru",
]
MIGRATIONS_SCOPE_DECISION: list[str] = []
MIGRATIONS_EXCLUDED = ["elektro.spb.ru"]
MIGRATIONS_ADDITIONAL = ["39mchs.ru"]

MIGRATION_EVIDENCE_DOMAINS = MIGRATIONS_LIVE + MIGRATIONS_SOURCE_PLACEHOLDER + MIGRATIONS_ADDITIONAL
MIGRATION_SPECIAL_QA = {
    "91web.ru",
    "moopb.ru",
    "electro-reg.ru",
    "othodi-spb.ru",
    "rectavr.ru",
    "fstek.spb.ru",
}

FORM_REQUIREMENTS = [
    ("Две формы", "На каждом из 30 включённых сайтов есть «ЗАКАЗАТЬ ЗВОНОК» и «ЗАДАТЬ ВОПРОС»."),
    ("Названия", "Заголовки и кнопки приведены к одинаковым согласованным названиям."),
    ("Форма звонка", "Имя необязательно; телефон обязателен; email и выбор типа клиента убраны; капча есть."),
    ("Форма вопроса", "Имя и вопрос необязательны; телефон обязателен; email убран; капча есть."),
    ("Согласие", "Размещён согласованный текст со ссылкой на политику; отдельной галочки нет."),
    ("Закрытие окна", "Крестик виден, окно закрывается и не перекрывается чатом или меню."),
    ("Ошибочные сценарии", "Убраны пустые страницы, 404, вечная загрузка и ложное сообщение об успехе."),
    ("Успешная отправка", "Подтверждение показывается только после фактического принятия заявки."),
    ("Обработчики", "Обе формы на каждом включённом сайте приняли тестовую заявку."),
    (
        "Почта",
        "Получатели всех форм сверены со штатной матрицей. Фактическое получение обеих форм отдельно подтверждено в доступном ящике info@medlic.spb.ru; по остальным сайтам заявлено только принятие обработчиком.",
    ),
    ("Внешний вид", "После публикации каждый сайт просмотрен на компьютере и телефоне."),
]

CLIENT_INPUT_REQUIRED = [
    (
        "dpocenter.ru",
        "Сайт работает на Sprinthost. В письмах и переданных файлах нет доступа к этому хостингу и нет полной копии сайта.",
        "Доступ к Sprinthost или свежая резервная копия файлов и базы.",
    ),
    (
        "feo-edem.ru",
        "Найдена база данных, но файлов сайта нет. Сам домен сейчас не зарегистрирован.",
        "Файлы сайта или доступ к старому хостингу, а также зарегистрированный домен.",
    ),
    (
        "linkedin.com.moopb.ru",
        "Поддомен не работает. В старом источнике была только страница-заглушка, копии настоящего сайта нет.",
        "Подтвердить, нужен ли этот поддомен. Если нужен - прислать исходники или пример нужного содержимого.",
    ),
    (
        "mchs-vrn.ru",
        "Копия сайта подготовлена на Beget и проверена, но домен сейчас не зарегистрирован.",
        "Зарегистрировать или продлить домен и направить его на Beget.",
    ),
    (
        "aklab-spb.ru",
        "Домен не зарегистрирован. Файлов сайта и доступа к старому хостингу в переданных материалах нет.",
        "Зарегистрированный домен и исходники сайта либо доступ к прежнему хостингу.",
    ),
    (
        "othodi-spb.ru",
        "Скопировано всё, что было в доступном источнике, но там оказалась только стандартная страница хостинга.",
        "Файлы настоящего сайта или подтверждение, что вместо сайта и раньше должна была быть эта страница.",
    ),
]

COVER_NOTE_PARAGRAPHS = [
    "Альберт, добрый день.",
    (
        "В прошлый раз я поспешил с отчётом: проверил не все сайты и написал, что работа закончена. "
        "Вы справедливо указали, что часть замечаний осталась. Это моя ошибка."
    ),
    (
        "После вашего письма я ещё раз поднял все материалы, прошёл каждый сайт из последней доработки "
        "на компьютере и телефоне и исправил то, что пропустил. Новый список замечаний от вас не нужен."
    ),
    (
        "В приложенном отчёте по каждому пункту прямо написано, что вы просили, что я сделал неправильно "
        "и что исправил. В конце оставил только шесть старых позиций по переносу, которые нельзя закрыть "
        "без отсутствующего домена или исходников. Все доступы, которые вы уже присылали, я учёл и повторно не запрашиваю."
    ),
    "С уважением,\nНикита Тихомиров",
]

HUMAN_CORRECTIONS = [
    (
        "Разные типы сайтов",
        "Формы нужно было проверить отдельно на сайтах «АП-Риал» и на сайтах лицензирования.",
        "Я использовал один подход для разных макетов и не проверил глазами результат на каждом сайте.",
        "Разделил сайты по типам и проверил каждый из них отдельно. Формы сохранили оформление своего сайта и не ломают страницу.",
    ),
    (
        "Поля в формах",
        "В обратном звонке должны быть имя, телефон и капча. В форме вопроса - имя, телефон, необязательный вопрос и капча.",
        "На части сайтов оставались лишние поля, а вопрос был обязательным.",
        "Лишние поля убраны. Имя и текст вопроса можно не заполнять; телефон остался обязательным. Капча есть в обеих формах.",
    ),
    (
        "Вид на телефоне",
        "Окно должно целиком помещаться на экране, а крестик - быть видимым и работать.",
        "На некоторых сайтах окно обрезалось или перекрывалось меню и чатом.",
        "Исправлены размеры и положение окон. Каждая форма заново открыта и закрыта на компьютере и телефоне.",
    ),
    (
        "Отправка заявки",
        "После отправки посетитель должен увидеть понятный результат, а письмо - уйти в рабочую почту сайта.",
        "На части сайтов после отправки появлялась пустая страница, бесконечная загрузка или сообщение об успехе до приёма заявки.",
        "Исправил обработчики. Теперь сообщение об успехе появляется только после того, как сайт принял заявку.",
    ),
    (
        "Адреса получателей",
        "Каждый сайт должен отправлять заявки на свой рабочий адрес, а не в один общий ящик.",
        "В прошлом отчёте я показал работу форм, но не показал, какие адреса стоят в настройках. Поэтому было непонятно, куда уходят письма.",
        "Заново проверил адрес получателя в каждой форме. Личный Gmail нигде не указан: каждый сайт отправляет заявки на свой рабочий адрес.",
    ),
    (
        "Фоновое видео",
        "Старый видеофон на nousro.ru и nousro-nn.ru не входил в текущую задачу.",
        "При исправлении ошибок JavaScript этот фон снова стал видимым, хотя менять его не просили.",
        "Видеофон снова скрыт на обоих сайтах. Остальные исправления JavaScript сохранены.",
    ),
]

ADDITIONAL_WORK = [
    (
        "mchs-spb.ru",
        "Почта и формы",
        "Проверены DNS, MX/SPF/DKIM, обработчики и фактическое получение обеих форм.",
    ),
    (
        "medlic.spb.ru",
        "Видимая часть и SEO",
        "Опубликованы согласованные изменения; слайдер и страницы проверены на компьютере и телефоне.",
    ),
    (
        "apreal.ru",
        "Доступ WordPress",
        "Проверены рабочие адреса входа в панель управления.",
    ),
    (
        "ohrana-truda.nousro.ru",
        "Восстановление",
        "Сайт восстановлен, старые ошибки ресурсов и сценариев устранены, desktop/mobile проверены заново.",
    ),
    (
        "rectavr.ru",
        "Мобильная версия",
        "Во время финального контроля найдено и исправлено горизонтальное обрезание шапки на телефоне.",
    ),
    (
        "mchs-vrn.ru",
        "Подготовленная версия",
        "Исправлено мобильное обрезание баннера; версия на Beget повторно проверена до подключения домена.",
    ),
    (
        "30 сайтов с формами",
        "JavaScript и ресурсы",
        "После финальной публикации свежий прогон не выявил ошибок страницы, критических ошибок консоли или сорванных запросов.",
    ),
    (
        "nousro.ru / nousro-nn.ru",
        "Фоновый видеоэлемент",
        "Старый фоновый видеоэлемент временно появился при устранении JavaScript-ошибок, но не относился к поручению. После проверки он скрыт на обоих сайтах.",
    ),
]


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def load_result_index(path: Path = FRESH_QA_RESULTS) -> dict[tuple[str, str], dict[str, Any]]:
    index: dict[tuple[str, str], dict[str, Any]] = {}
    for item in load_json(path):
        if item.get("type") == "excluded":
            continue
        key = (item["domain"], item["viewport"])
        if key in index:
            raise ValueError(f"Duplicate QA result: {key[0]} {key[1]}")
        index[key] = item
    return index


def load_delivery_index() -> dict[tuple[str, str], dict[str, Any]]:
    submissions = load_json(SENDER_DELIVERY_PATH)["submissions"]
    return {(item["domain"], item["kind"]): item for item in submissions}


def as_path(raw: str | Path) -> Path:
    path = Path(raw)
    return path if path.is_absolute() else ROOT / path


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = "Calibri"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = MID_GRAY, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_table_width(table, widths: list[float]) -> None:
    evidence.set_table_width(table, widths)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def add_field(paragraph, instruction: str) -> None:
    evidence.add_field(paragraph, instruction)


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def configure_document(doc: Document, *, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run(title), size=8.5, bold=True, color=DARK_GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    run = footer.add_run("Страница ")
    set_run_font(run, size=8.5, color=DARK_GRAY)
    add_field(footer, "PAGE")
    run = footer.add_run(" из ")
    set_run_font(run, size=8.5, color=DARK_GRAY)
    add_field(footer, "NUMPAGES")

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_metadata_table(doc: Document, rows: Iterable[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    set_table_width(table, [1.45, 5.05])
    set_table_borders(table)
    for label, value in rows:
        row = table.add_row()
        keep_row_together(row)
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label), size=9, bold=True, color=DARK_BLUE)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), size=9.5, color=BLACK)


def add_metric_strip(doc: Document, metrics: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(metrics))
    set_table_width(table, [6.5 / len(metrics)] * len(metrics))
    set_table_borders(table, color="B8D4E8", size=7)
    for cell, (value, label) in zip(table.rows[0].cells, metrics):
        set_cell_shading(cell, PALE_BLUE)
        set_cell_margins(cell, top=115, bottom=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        set_run_font(p.add_run(value), size=15, bold=True, color=BLUE)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label), size=7.8, color=DARK_GRAY)


def add_callout(doc: Document, text: str, *, status: str = "good") -> None:
    styles = {
        "good": (PALE_GREEN, "B9D7C5", GREEN),
        "note": (PALE_BLUE, "B8D4E8", DARK_BLUE),
        "warn": (PALE_AMBER, "E7C77D", AMBER),
    }
    fill, border, color = styles[status]
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    set_table_borders(table, color=border, size=7)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=125, bottom=125, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(text), size=10, bold=True, color=color)


def add_picture_with_caption(cell, path: Path, caption: str, width: float) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell, top=70, bottom=70)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width))
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(caption), size=8, color=DARK_GRAY)


def add_report_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run("ИТОГОВЫЙ ОТЧЁТ"), size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run("Работы по сайтам ГК «АП-Риал»"), size=24, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    set_run_font(
        p.add_run("Перенос сайтов, формы обратной связи, маршрутизация заявок и последующие исправления"),
        size=13,
        color=DARK_GRAY,
    )

    add_metadata_table(
        doc,
        [
            ("Заказчик", "Группа компаний «АП-Риал»"),
            ("Исполнитель", "Никита Тихомиров"),
            ("Дата", "2 августа 2026 года"),
            ("Проверка", "Повторная сверка исходных поручений и опубликованных версий"),
        ],
    )

    doc.add_heading("Краткий итог", level=2)
    add_metric_strip(
        doc,
        [
            ("30", "сайтов с формами"),
            ("60 из 60", "заявок приняты"),
            ("48 из 48", "маршрутов верны"),
            ("28", "переносов в работе"),
        ],
    )

    p = doc.add_paragraph(
        "Все требования к формам на 30 включённых сайтах выполнены и повторно проверены после публикации. "
        "По переносу 28 сайтов работают с перенесёнными версиями; один доступный источник оказался "
        "стандартной страницей хостинга. По остальным позициям ниже указано только то, что объективно "
        "невозможно завершить без домена, исходников, доступа или решения заказчика."
    )
    p.paragraph_format.space_before = Pt(12)
    add_callout(
        doc,
        "Повторно составлять список замечаний не требуется: исходные поручения восстановлены и сверены заново.",
    )


def add_form_results(doc: Document) -> None:
    add_page_break(doc)
    doc.add_heading("1. Формы обратной связи", level=1)
    p = doc.add_paragraph(
        "На 30 сайтах выполнена единая доработка двух форм. Внешний вид и работа обеих форм "
        "проверены на каждом сайте; все контрольные заявки приняты обработчиками. Получатели "
        "сверены отдельно по полной матрице маршрутов."
    )
    p.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [1.6, 4.0, 0.9])
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ("Что требовалось", "Что сделано", "Статус")):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=8.5, bold=True, color=DARK_BLUE)
    evidence.set_repeat_table_header(table.rows[0])
    for title, detail in FORM_REQUIREMENTS:
        row = table.add_row()
        keep_row_together(row)
        for cell in row.cells:
            set_cell_margins(cell, top=62, bottom=62)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(title), size=8.4, bold=True, color=BLACK)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(detail), size=8.3, color=DARK_GRAY)
        set_cell_shading(row.cells[2], PALE_GREEN)
        p = row.cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run("Готово"), size=8.3, bold=True, color=GREEN)

    doc.add_heading("Сайты, входившие в доработку", level=2)
    table = doc.add_table(rows=1, cols=4)
    set_table_width(table, [2.25, 1.0, 2.25, 1.0])
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ("Сайт", "Статус", "Сайт", "Статус")):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=8.5, bold=True, color=DARK_BLUE)
    midpoint = len(evidence.INCLUDED_DOMAINS) // 2
    for left, right in zip(evidence.INCLUDED_DOMAINS[:midpoint], evidence.INCLUDED_DOMAINS[midpoint:]):
        row = table.add_row()
        keep_row_together(row)
        for cell, value, status_cell in zip(row.cells, (left, "Готово", right, "Готово"), (False, True, False, True)):
            set_cell_margins(cell, top=48, bottom=48)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if status_cell:
                set_cell_shading(cell, PALE_GREEN)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(p.add_run(value), size=8.1, bold=True, color=GREEN)
            else:
                set_run_font(p.add_run(value), size=8.3, bold=True, color=BLACK)

    add_page_break(doc)
    doc.add_heading("Сайты, где формы не требовались", level=2)
    p = doc.add_paragraph(
        "Формы не добавлялись на пяти сайтах по прямому указанию. Это не пропуск в работе."
    )
    p.paragraph_format.space_after = Pt(5)
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [2.2, 4.3])
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ("Сайт", "Статус")):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=8.5, bold=True, color=DARK_BLUE)
    for domain in evidence.EXCLUDED_DOMAINS:
        row = table.add_row()
        keep_row_together(row)
        for cell in row.cells:
            set_cell_margins(cell, top=48, bottom=48)
        set_run_font(row.cells[0].paragraphs[0].add_run(domain), size=8.3, bold=True)
        set_run_font(
            row.cells[1].paragraphs[0].add_run("Формы исключены из согласованного объёма"),
            size=8.3,
            color=DARK_GRAY,
        )


def add_delivery_results(doc: Document) -> None:
    add_page_break(doc)
    doc.add_heading("2. Отправка и маршрутизация заявок", level=1)
    p = doc.add_paragraph(
        "Выполнено 60 контрольных отправок: по две формы на каждом из 30 сайтов. "
        "Все 60 обработчиков приняли заявки. Отдельная проверка 48 актуальных и дополнительных "
        "конфигураций показала, что каждая форма направляет письмо на адрес из клиентской матрицы. "
        "Фактическое получение обеих форм в почте отдельно подтверждено для medlic.spb.ru."
    )
    p.paragraph_format.space_after = Pt(8)
    add_metric_strip(
        doc,
        [("60", "отправлено"), ("60", "принято"), ("48", "маршрутов верны"), ("2", "письма в medlic")],
    )

    doc.add_heading("Подтверждённое получение в почте", level=2)
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    set_table_borders(table)
    add_picture_with_caption(
        table.cell(0, 0),
        MAILBOX_RECEIPT_SCREENSHOT,
        "В доступном ящике info@medlic.spb.ru подтверждены письма от обеих форм сайта medlic.spb.ru.",
        6.22,
    )
    add_callout(
        doc,
        "Для остальных сайтов подтверждены принятие заявки обработчиком и правильный адрес "
        "получателя в конфигурации. Получение письма непосредственно в каждом закрытом клиентском "
        "ящике без доступа к этому ящику в отчёте не заявляется.",
    )


def add_domain_list_table(doc: Document, domains: list[str], *, label: str) -> None:
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [2.16, 2.17, 2.17])
    set_table_borders(table)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label), size=8.5, bold=True, color=DARK_BLUE)
    rows = (len(domains) + 2) // 3
    columns = [domains[i * rows : (i + 1) * rows] for i in range(3)]
    for index in range(rows):
        row = table.add_row()
        keep_row_together(row)
        for col, cell in enumerate(row.cells):
            set_cell_margins(cell, top=44, bottom=44)
            value = columns[col][index] if index < len(columns[col]) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(value), size=8.2, bold=bool(value), color=BLACK)


def add_migration_results(doc: Document) -> None:
    add_page_break(doc)
    doc.add_heading("3. Перенос сайтов на Beget", level=1)
    p = doc.add_paragraph(
        "Исходное поручение охватывало 35 доменов. Ниже статус разделён так, чтобы перенесённый "
        "рабочий сайт не смешивался с заглушкой, подготовленной версией без домена или отсутствующим источником."
    )
    p.paragraph_format.space_after = Pt(8)

    add_callout(
        doc,
        "28 прямо порученных сайтов работают с перенесёнными версиями. Дополнительно перенесён 39mchs.ru.",
    )
    doc.add_heading("Работающие перенесённые версии", level=2)
    add_domain_list_table(doc, MIGRATIONS_LIVE, label="Перенесено и открывается")

    doc.add_heading("Особые позиции", level=2)
    rows = [
        (
            "othodi-spb.ru",
            "Перенесён доступный источник",
            "В источнике была только стандартная страница хостинга; за рабочий сайт это не выдаётся.",
            "Требуется источник",
        ),
        (
            "mchs-vrn.ru",
            "Подготовлено на Beget",
            "Файлы и база размещены, desktop/mobile работают; публичный домен недоступен.",
            "Нужны домен и DNS",
        ),
        (
            "39mchs.ru",
            "Перенесено дополнительно",
            "Сайт работает и входит в текущую проверку форм.",
            "Готово",
        ),
        (
            "elektro.spb.ru",
            "Не завершено",
            "Есть противоречие по объёму: поручение найдено, подтверждённого исключения нет.",
            "Нужно решение",
        ),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_width(table, [1.3, 1.45, 2.7, 1.05])
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ("Сайт", "Состояние", "Что это означает", "Итог")):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=8.2, bold=True, color=DARK_BLUE)
    for domain, state, meaning, result in rows:
        row = table.add_row()
        keep_row_together(row)
        for cell in row.cells:
            set_cell_margins(cell, top=58, bottom=58)
        values = (domain, state, meaning, result)
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if index == 0:
                set_run_font(p.add_run(value), size=8.1, bold=True)
            elif index == 3:
                fill = PALE_GREEN if value == "Готово" else PALE_AMBER
                color = GREEN if value == "Готово" else AMBER
                set_cell_shading(cell, fill)
                set_run_font(p.add_run(value), size=8.0, bold=True, color=color)
            else:
                set_run_font(p.add_run(value), size=8.0, color=DARK_GRAY)


def add_additional_results(doc: Document) -> None:
    add_page_break(doc)
    doc.add_heading("4. Последующие исправления", level=1)
    p = doc.add_paragraph(
        "Отдельные поручения и дефекты не смешивались с массовой доработкой форм. "
        "Каждый пункт ниже проверен отдельно."
    )
    p.paragraph_format.space_after = Pt(8)
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [1.6, 1.6, 3.3])
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ("Сайт", "Работа", "Фактический результат")):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=8.5, bold=True, color=DARK_BLUE)
    evidence.set_repeat_table_header(table.rows[0])
    for domain, work, result in ADDITIONAL_WORK:
        row = table.add_row()
        keep_row_together(row)
        for cell in row.cells:
            set_cell_margins(cell, top=64, bottom=64)
        set_run_font(row.cells[0].paragraphs[0].add_run(domain), size=8.2, bold=True)
        set_run_font(row.cells[1].paragraphs[0].add_run(work), size=8.2, bold=True, color=DARK_BLUE)
        set_run_font(row.cells[2].paragraphs[0].add_run(result), size=8.2, color=DARK_GRAY)

    doc.add_heading("Скрытый фоновый видеоэлемент и существующая камера", level=2)
    p = doc.add_paragraph(
        "Фоновый видеоэлемент на nousro.ru и nousro-nn.ru не входил в поручение. Он временно появился "
        "при устранении JavaScript-ошибок, после чего был повторно проверен и скрыт на обоих сайтах. "
        "Дополнительное решение заказчика по этому элементу не требуется."
    )
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph(
        "Блок Ivideon — другой, ранее существовавший элемент. Он не изменялся и сейчас не подключается к камере. "
        "Для восстановления нужен действующий доступ/идентификатор камеры; без него блок можно только скрыть по согласованию."
    )
    p.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [3.25, 3.25])
    set_table_borders(table)
    add_picture_with_caption(
        table.cell(0, 0),
        POST_CORRECTION_QA_DIR / "nousro.ru-desktop-page.png",
        "nousro.ru после исправления: фоновый видеоэлемент скрыт.",
        3.02,
    )
    add_picture_with_caption(
        table.cell(0, 1),
        POST_CORRECTION_QA_DIR / "nousro-nn.ru-desktop-page.png",
        "nousro-nn.ru после исправления: фоновый видеоэлемент скрыт.",
        3.02,
    )
    add_page_break(doc)
    doc.add_heading("Существующий блок Ivideon", level=2)
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [3.25, 3.25])
    set_table_borders(table)
    add_picture_with_caption(
        table.cell(0, 0),
        VIDEO_QA_DIR / "nousro.ru-existing-ivideon-block.png",
        "Существующий блок Ivideon: камера сейчас недоступна.",
        3.02,
    )
    add_picture_with_caption(
        table.cell(0, 1),
        VIDEO_QA_DIR / "nousro-nn.ru-existing-ivideon-block.png",
        "На втором сайте тот же отдельный блок камеры.",
        3.02,
    )


def add_client_inputs(doc: Document) -> None:
    add_page_break(doc)
    doc.add_heading("5. Что требуется от заказчика", level=1)
    p = doc.add_paragraph(
        "Ниже нет работ, которые можно было бы просто доделать без участия заказчика. "
        "Указаны только отсутствующие домены, источники, доступы или решения."
    )
    p.paragraph_format.space_after = Pt(8)
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [1.55, 2.65, 2.30])
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ("Объект", "Фактическое состояние", "Что именно нужно")):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=8.4, bold=True, color=DARK_BLUE)
    evidence.set_repeat_table_header(table.rows[0])
    for subject, state, needed in CLIENT_INPUT_REQUIRED:
        row = table.add_row()
        keep_row_together(row)
        for cell in row.cells:
            set_cell_margins(cell, top=64, bottom=64)
        set_run_font(row.cells[0].paragraphs[0].add_run(subject), size=8.1, bold=True)
        set_run_font(row.cells[1].paragraphs[0].add_run(state), size=8.0, color=DARK_GRAY)
        set_cell_shading(row.cells[2], PALE_AMBER)
        set_run_font(row.cells[2].paragraphs[0].add_run(needed), size=8.0, color=AMBER)

    add_callout(
        doc,
        "Все остальные перечисленные в исходных поручениях исправления выполнены. Открытые позиции выше не выданы за завершённые.",
        status="note",
    )


def add_special_evidence(doc: Document) -> None:
    add_page_break(doc)
    doc.add_heading("6. Примеры отдельных исправлений", level=1)

    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [3.25, 3.25])
    set_table_borders(table)
    add_picture_with_caption(
        table.cell(0, 0),
        MAIL_EVIDENCE_DIR / "medlic.spb.ru-slider-desktop.png",
        "medlic.spb.ru: опубликованная версия на компьютере.",
        3.02,
    )
    add_picture_with_caption(
        table.cell(0, 1),
        MAIL_EVIDENCE_DIR / "medlic.spb.ru-slider-mobile.png",
        "medlic.spb.ru: та же версия на телефоне.",
        3.02,
    )

    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [3.25, 3.25])
    set_table_borders(table)
    add_picture_with_caption(
        table.cell(0, 0),
        MIGRATION_QA_DIR / "mchs-vrn.ru-desktop-staged.png",
        "mchs-vrn.ru: версия на Beget до подключения публичного домена.",
        3.02,
    )
    add_picture_with_caption(
        table.cell(0, 1),
        MIGRATION_QA_DIR / "mchs-vrn.ru-mobile-staged.png",
        "mchs-vrn.ru: исправленная мобильная версия.",
        3.02,
    )

    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [3.25, 3.25])
    set_table_borders(table)
    add_picture_with_caption(
        table.cell(0, 0),
        MIGRATION_QA_DIR / "ohrana-truda.nousro.ru-desktop-final-acceptance.png",
        "ohrana-truda.nousro.ru: восстановленная desktop-версия.",
        3.02,
    )
    add_picture_with_caption(
        table.cell(0, 1),
        MIGRATION_QA_DIR / "ohrana-truda.nousro.ru-mobile-final-acceptance.png",
        "ohrana-truda.nousro.ru: восстановленная mobile-версия.",
        3.02,
    )


def build_form_board(domain: str, result_index: dict[tuple[str, str], dict[str, Any]]) -> Path:
    target = ASSET_DIR / f"{domain}-forms.png"
    canvas = Image.new("RGB", (1800, 1040), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = evidence.get_font(38, bold=True)
    label_font = evidence.get_font(23, bold=True)
    small_font = evidence.get_font(20)
    status_font = evidence.get_font(22, bold=True)

    draw.text((50, 26), domain, font=title_font, fill="#1F4D78")
    draw.text((50, 75), "Обе формы после публикации: компьютер и телефон", font=small_font, fill="#4B5563")
    status = "ПРОВЕРЕНО"
    status_box = draw.textbbox((0, 0), status, font=status_font)
    draw.text((1750 - (status_box[2] - status_box[0]), 45), status, font=status_font, fill="#1F7A45")

    panels = [
        ("desktop", "callback", "КОМПЬЮТЕР · ЗАКАЗАТЬ ЗВОНОК"),
        ("desktop", "question", "КОМПЬЮТЕР · ЗАДАТЬ ВОПРОС"),
        ("mobile", "callback", "ТЕЛЕФОН · ЗАКАЗАТЬ ЗВОНОК"),
        ("mobile", "question", "ТЕЛЕФОН · ЗАДАТЬ ВОПРОС"),
    ]
    boxes = [(45, 150, 885, 565), (915, 150, 1755, 565), (45, 610, 885, 1025), (915, 610, 1755, 1025)]
    for (viewport, kind, label), box in zip(panels, boxes):
        result = result_index[(domain, viewport)]
        screenshot = evidence.screenshot_for(result, kind)
        rect = result.get("actions", {}).get(kind, {}).get("modal", {}).get("rect")
        cropped = evidence.crop_modal(screenshot, rect)
        left, top, right, bottom = box
        draw.rounded_rectangle(box, radius=8, outline="#B9C6D3", width=2, fill="#F8FAFC")
        draw.text((left + 16, top + 12), label, font=label_font, fill="#111111")
        evidence.paste_contained(canvas, cropped, (left + 12, top + 52, right - 12, bottom - 12))

    target.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(target, optimize=True)
    return target


def migration_screenshot(domain: str, viewport: str) -> Path:
    if domain in MIGRATION_SPECIAL_QA:
        return MIGRATION_QA_DIR / f"{domain}-{viewport}-final-acceptance.png"
    return FRESH_QA_DIR / f"{domain}-{viewport}-page.png"


def build_migration_sheet(domains: list[str], index: int) -> Path:
    target = ASSET_DIR / f"migration-sheet-{index:02d}.png"
    canvas = Image.new("RGB", (1800, 2200), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = evidence.get_font(34, bold=True)
    label_font = evidence.get_font(23, bold=True)
    small_font = evidence.get_font(18)
    draw.text((55, 28), "Перенесённые версии: свежая desktop/mobile проверка", font=title_font, fill="#1F4D78")

    row_height = 505
    for row_index, domain in enumerate(domains):
        top = 105 + row_index * row_height
        draw.text((55, top), domain, font=label_font, fill="#111111")
        draw.text((1540, top + 2), "ПРОВЕРЕНО", font=small_font, fill="#1F7A45")
        desktop_box = (55, top + 42, 1240, top + 470)
        mobile_box = (1270, top + 42, 1745, top + 470)
        for viewport, box in (("desktop", desktop_box), ("mobile", mobile_box)):
            screenshot = migration_screenshot(domain, viewport)
            if not screenshot.exists():
                raise FileNotFoundError(screenshot)
            image = Image.open(screenshot).convert("RGB")
            draw.rounded_rectangle(box, radius=8, outline="#B9C6D3", width=2, fill="#F8FAFC")
            evidence.paste_contained(canvas, image, (box[0] + 8, box[1] + 8, box[2] - 8, box[3] - 8))

    target.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(target, optimize=True)
    return target


def add_form_appendix(doc: Document, boards: dict[str, Path]) -> None:
    for page_index in range(0, len(evidence.INCLUDED_DOMAINS), 2):
        add_page_break(doc)
        doc.add_heading("Приложение 1. Формы по каждому сайту", level=1)
        for domain in evidence.INCLUDED_DOMAINS[page_index : page_index + 2]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            set_run_font(p.add_run(domain), size=10, bold=True, color=DARK_BLUE)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            p.add_run().add_picture(str(boards[domain]), width=Inches(6.2))


def add_migration_appendix(doc: Document, sheets: list[Path]) -> None:
    for sheet in sheets:
        add_page_break(doc)
        doc.add_heading("Приложение 2. Перенесённые сайты", level=1)
        p = doc.add_paragraph(
            "На снимках показаны опубликованные версии после повторной проверки на компьютере и телефоне."
        )
        p.paragraph_format.space_after = Pt(4)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(sheet), width=Inches(6.2))


def crop_report_image(source: Path, target: Path, box: tuple[int, int, int, int]) -> Path:
    if not source.exists():
        raise FileNotFoundError(source)
    target.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(source) as image:
        left, top, right, bottom = box
        right = min(right, image.width)
        bottom = min(bottom, image.height)
        if left >= right or top >= bottom:
            raise ValueError(f"Invalid crop for {source}: {box}, image={image.size}")
        image.crop((left, top, right, bottom)).convert("RGB").save(
            target,
            format="PNG",
            optimize=True,
        )
    return target


def build_human_report_assets(result_index: dict[tuple[str, str], dict[str, Any]]) -> dict[str, Path]:
    def action_screenshot(domain: str, viewport: str, kind: str) -> Path:
        raw = result_index[(domain, viewport)]["actions"][kind]["screenshot"]
        return as_path(raw)

    assets = {
        "apreal_callback": crop_report_image(
            action_screenshot("apreal.ru", "desktop", "callback"),
            ASSET_DIR / "apreal.ru-callback-crop.png",
            (330, 20, 1110, 575),
        ),
        "apreal_question": crop_report_image(
            action_screenshot("apreal.ru", "desktop", "question"),
            ASSET_DIR / "apreal.ru-question-crop.png",
            (330, 20, 1110, 710),
        ),
        "license_callback": crop_report_image(
            action_screenshot("license39.ru", "desktop", "callback"),
            ASSET_DIR / "license39.ru-callback-crop.png",
            (415, 145, 1025, 860),
        ),
        "license_question": crop_report_image(
            action_screenshot("license39.ru", "desktop", "question"),
            ASSET_DIR / "license39.ru-question-crop.png",
            (415, 145, 1025, 860),
        ),
        "mobile_callback": crop_report_image(
            action_screenshot("apreal-nn.ru", "mobile", "callback"),
            ASSET_DIR / "apreal-nn.ru-mobile-callback.png",
            (0, 0, 390, 844),
        ),
        "mobile_question": crop_report_image(
            action_screenshot("nousro.ru", "mobile", "question"),
            ASSET_DIR / "nousro.ru-mobile-question.png",
            (0, 0, 390, 844),
        ),
        "mailbox": crop_report_image(
            MAILBOX_RECEIPT_SCREENSHOT,
            ASSET_DIR / "medlic-mailbox-crop.png",
            (470, 175, 2190, 535),
        ),
        "nousro_video": crop_report_image(
            ROOT / "output/ap-real-hidden-video-sections-2026-08-05/nousro.ru-video-section-hidden.png",
            ASSET_DIR / "nousro.ru-hidden-video-crop.png",
            (0, 500, 1440, 2050),
        ),
        "nousro_nn_video": crop_report_image(
            ROOT / "output/ap-real-hidden-video-sections-2026-08-05/nousro-nn.ru-video-section-hidden.png",
            ASSET_DIR / "nousro-nn.ru-hidden-video-crop.png",
            (0, 500, 1440, 2050),
        ),
    }
    return assets


def add_evidence_pair(
    doc: Document,
    left: tuple[Path, str],
    right: tuple[Path, str],
    *,
    width: float = 3.05,
) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [3.25, 3.25])
    set_table_borders(table, color=MID_GRAY, size=5)
    add_picture_with_caption(table.cell(0, 0), left[0], left[1], width)
    add_picture_with_caption(table.cell(0, 1), right[0], right[1], width)
    keep_row_together(table.rows[0])


def add_single_evidence(doc: Document, path: Path, caption: str, *, width: float = 6.1) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    set_table_borders(table, color=MID_GRAY, size=5)
    add_picture_with_caption(table.cell(0, 0), path, caption, width)
    keep_row_together(table.rows[0])


def add_human_report_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run("ОТЧЁТ ПОСЛЕ ПОВТОРНОЙ ПРОВЕРКИ"), size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("Что исправлено на сайтах «АП-Риал»"), size=24, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    set_run_font(
        p.add_run("Формы обратной связи, вид на телефоне, адреса заявок и фоновое видео"),
        size=13,
        color=DARK_GRAY,
    )

    add_metadata_table(
        doc,
        [
            ("Заказчик", "Группа компаний «АП-Риал»"),
            ("Исполнитель", "Никита Тихомиров"),
            ("Дата", "5 августа 2026 года"),
            ("О чём отчёт", "Последние замечания по формам и внешнему виду сайтов"),
        ],
    )

    doc.add_heading("Коротко", level=2)
    p = doc.add_paragraph(
        "В прошлый раз я проверил не все сайты и слишком рано написал, что работа закончена. После вашего письма "
        "я заново прошёл всю последнюю доработку: открыл формы на компьютере и телефоне, проверил отправку заявок "
        "и адреса получателей. Ниже показано, что было пропущено и что исправлено."
    )
    p.paragraph_format.space_after = Pt(12)
    add_callout(doc, "Повторно составлять список замечаний не нужно. Ниже указано, что именно было исправлено.")


def add_human_corrections(doc: Document) -> None:
    add_page_break(doc)
    doc.add_heading("Что было не так и что исправлено", level=1)
    p = doc.add_paragraph(
        "Ниже без общих формулировок: что вы просили, что я сделал неправильно и что исправил."
    )
    p.paragraph_format.space_after = Pt(8)

    for index, (title, required, wrong, fixed) in enumerate(HUMAN_CORRECTIONS, start=1):
        table = doc.add_table(rows=1, cols=1)
        set_table_width(table, [6.5])
        set_table_borders(table, color="B8D4E8", size=6)
        cell = table.cell(0, 0)
        set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
        if index % 2:
            set_cell_shading(cell, "F8FAFC")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(5)
        set_run_font(p.add_run(f"{index}. {title}"), size=11.5, bold=True, color=DARK_BLUE)
        for label, text, color in (
            ("Нужно было", required, DARK_BLUE),
            ("Что было не так", wrong, RED),
            ("Что исправлено", fixed, GREEN),
        ):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            set_run_font(p.add_run(f"{label}: "), size=9.2, bold=True, color=color)
            set_run_font(p.add_run(text), size=9.2, color=BLACK)
        keep_row_together(table.rows[0])
        doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_human_visual_evidence(doc: Document, assets: dict[str, Path]) -> None:
    add_page_break(doc)
    doc.add_heading("Как формы выглядят сейчас", level=1)
    p = doc.add_paragraph(
        "Ниже - свежие снимки после исправлений. Показываю два разных типа сайтов: у каждого осталось своё оформление, изменились только поля и работа форм."
    )
    doc.add_heading("Сайт «АП-Риал»", level=2)
    add_evidence_pair(
        doc,
        (assets["apreal_callback"], "apreal.ru - обратный звонок: имя, телефон и капча."),
        (assets["apreal_question"], "apreal.ru - вопрос: текст вопроса можно не заполнять."),
    )

    add_page_break(doc)
    doc.add_heading("Сайт лицензирования", level=1)
    p = doc.add_paragraph(
        "У license39.ru другое оформление. Я его не менял и привёл в порядок только поля и работу форм."
    )
    add_evidence_pair(
        doc,
        (assets["license_callback"], "license39.ru - форма «Заказать звонок»."),
        (assets["license_question"], "license39.ru - форма «Задать вопрос»."),
    )

    add_page_break(doc)
    doc.add_heading("Проверка на телефоне", level=1)
    p = doc.add_paragraph(
        "Окна помещаются на экране, поля не обрезаны, крестик закрытия виден. На снимках - два сайта с разной вёрсткой."
    )
    add_evidence_pair(
        doc,
        (assets["mobile_callback"], "apreal-nn.ru - обратный звонок на телефоне."),
        (assets["mobile_question"], "nousro.ru - форма вопроса на телефоне."),
        width=2.55,
    )


def add_human_mail_and_video(doc: Document, assets: dict[str, Path]) -> None:
    add_page_break(doc)
    doc.add_heading("Куда приходят заявки", level=1)
    p = doc.add_paragraph(
        "Один общий адрес на все сайты не ставился. Я заново проверил адрес получателя в каждой рабочей форме. "
        "Личного Gmail в настройках нет: заявки уходят на рабочий адрес конкретного сайта."
    )
    p.paragraph_format.space_after = Pt(8)
    add_single_evidence(
        doc,
        assets["mailbox"],
        "Пример прямой проверки: в info@medlic.spb.ru видны письма от обеих форм medlic.spb.ru.",
    )
    add_callout(
        doc,
        "На остальных сайтах проверены адрес получателя в настройках формы и ответ сайта после отправки. Прямую доставку в ящик отдельно показываю только для medlic.spb.ru.",
        status="note",
    )

    add_page_break(doc)
    doc.add_heading("Фоновое видео снова скрыто", level=1)
    p = doc.add_paragraph(
        "При исправлении ошибок JavaScript на nousro.ru и nousro-nn.ru снова появился старый фон с движущимися цветными шарами. Он не относился к поручению и выглядел как ненужное изменение. "
        "Сейчас фон снова скрыт на обоих сайтах, а остальные исправления скриптов остались."
    )
    add_evidence_pair(
        doc,
        (assets["nousro_video"], "nousro.ru - раздел без старого видеофона."),
        (assets["nousro_nn_video"], "nousro-nn.ru - тот же раздел без видеофона."),
    )


def add_human_domain_list(doc: Document) -> None:
    add_page_break(doc)
    doc.add_heading("Какие сайты проверены", level=1)
    p = doc.add_paragraph(
        "Ниже полный список сайтов из последней доработки. На каждом открыты обе формы на компьютере и телефоне. Ошибок, обрезанных окон или нерабочих кнопок в этом списке не осталось."
    )
    domains = evidence.INCLUDED_DOMAINS
    columns = 3
    rows = (len(domains) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    set_table_width(table, [6.5 / columns] * columns)
    set_table_borders(table)
    for index, domain in enumerate(domains):
        row = index % rows
        column = index // rows
        cell = table.cell(row, column)
        set_cell_margins(cell, top=70, bottom=70)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(domain), size=9, bold=True, color=DARK_BLUE)
    for row in table.rows:
        keep_row_together(row)

    doc.add_heading("Сайты, где новые формы не требовались", level=2)
    p = doc.add_paragraph(
        "По вашему указанию новые формы на rectavr.ru, fstek.spb.ru, lic-k.ru, apreal-samara.ru и ed-krd.ru не делались. Поэтому в этом отчёте я не называю эти сайты исправленными."
    )
    p.paragraph_format.space_after = Pt(0)


def add_human_open_items(doc: Document) -> None:
    add_page_break(doc)
    doc.add_heading("Что осталось по старому переносу", level=1)
    p = doc.add_paragraph(
        "Эти шесть позиций не связаны с последними замечаниями по формам. Для них действительно не хватает домена или исходников. Уже присланные доступы здесь повторно не запрашиваются."
    )
    p.paragraph_format.space_after = Pt(8)

    for index, (domain, state, needed) in enumerate(CLIENT_INPUT_REQUIRED, start=1):
        if index == 4:
            add_page_break(doc)
        table = doc.add_table(rows=1, cols=1)
        set_table_width(table, [6.5])
        set_table_borders(table, color="E7C77D", size=6)
        cell = table.cell(0, 0)
        set_cell_shading(cell, PALE_AMBER)
        set_cell_margins(cell, top=105, bottom=105, start=145, end=145)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run(domain), size=11, bold=True, color=DARK_BLUE)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run("Сейчас: "), size=9.2, bold=True, color=DARK_GRAY)
        set_run_font(p.add_run(state), size=9.2, color=BLACK)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run("Нужно: "), size=9.2, bold=True, color=AMBER)
        set_run_font(p.add_run(needed), size=9.2, color=BLACK)
        keep_row_together(table.rows[0])
        doc.add_paragraph().paragraph_format.space_after = Pt(1)

    add_callout(
        doc,
        "По формам и сайтам из основной части отчёта дополнительных материалов от вас не требуется.",
    )


def validate_inputs(
    result_index: dict[tuple[str, str], dict[str, Any]],
    delivery_index: dict[tuple[str, str], dict[str, Any]],
) -> dict[str, Any]:
    expected_views = {(domain, viewport) for domain in evidence.INCLUDED_DOMAINS for viewport in ("desktop", "mobile")}
    if set(result_index) != expected_views:
        missing = sorted(expected_views - set(result_index))
        extra = sorted(set(result_index) - expected_views)
        raise ValueError(f"Fresh QA matrix mismatch. Missing={missing}, extra={extra}")

    screenshot_count = 0
    for key, result in result_index.items():
        problems = []
        if result.get("status") != 200:
            problems.append(f"HTTP {result.get('status')}")
        for field in ("failures", "pageErrors", "criticalConsoleErrors", "requestFailures"):
            if result.get(field):
                problems.append(f"{field}={result[field]}")
        if problems:
            raise ValueError(f"Fresh QA failed for {key[0]} {key[1]}: {'; '.join(problems)}")
        for kind in ("page", "callback", "question"):
            raw = result.get("screenshot") if kind == "page" else result.get("actions", {}).get(kind, {}).get("screenshot")
            if kind == "page" and not raw:
                raw = str(FRESH_QA_DIR / f"{key[0]}-{key[1]}-page.png")
            if not raw or not as_path(raw).exists():
                raise FileNotFoundError(f"Missing fresh screenshot: {key[0]} {key[1]} {kind}")
            screenshot_count += 1

    expected_delivery = {(domain, kind) for domain in evidence.INCLUDED_DOMAINS for kind in ("callback", "question")}
    if set(delivery_index) != expected_delivery:
        raise ValueError("Delivery matrix does not contain exactly 60 current domain/form pairs")
    for key, item in delivery_index.items():
        if not item.get("accepted"):
            raise ValueError(f"Submission was not accepted: {key[0]} {key[1]}")

    recipient_matrix = load_json(RECIPIENT_MATRIX_PATH)
    recipient_summary = recipient_matrix.get("summary", {})
    if (
        recipient_summary.get("checks") != 48
        or recipient_summary.get("passed") != 48
        or recipient_summary.get("failed")
        or recipient_summary.get("personal_recipient_hits")
        or recipient_summary.get("complete") is not True
    ):
        raise ValueError("Recipient matrix does not prove 48 correct production routes")

    mailbox_receipt = load_json(MAILBOX_RECEIPT_PATH)
    visible_accounts = set(mailbox_receipt.get("account_emails_visible_in_menu", []))
    if "info@medlic.spb.ru" not in visible_accounts:
        raise ValueError("Mailbox evidence is not tied to info@medlic.spb.ru")
    marker_hits = set(mailbox_receipt.get("marker_hits", []))
    if marker_hits != MAILBOX_RECEIPT_MARKERS:
        raise ValueError("Mailbox evidence does not prove both medlic.spb.ru form messages")
    if not MAILBOX_RECEIPT_SCREENSHOT.exists():
        raise FileNotFoundError(MAILBOX_RECEIPT_SCREENSHOT)

    hidden_video_evidence = load_json(HIDDEN_VIDEO_EVIDENCE_PATH)
    hidden_video_summary = hidden_video_evidence.get("summary", {})
    if (
        hidden_video_summary.get("checks") != 2
        or hidden_video_summary.get("passed") != 2
        or hidden_video_summary.get("failed")
        or hidden_video_summary.get("complete") is not True
    ):
        raise ValueError("Live evidence does not prove both background videos are hidden")

    migration_results = load_json(MIGRATION_QA_RESULTS)
    if len(migration_results) != 16:
        raise ValueError(f"Expected 16 migration QA views, got {len(migration_results)}")
    for item in migration_results:
        if item.get("status") != 200 or item.get("failures") or item.get("pageErrors") or item.get("criticalConsoleErrors"):
            raise ValueError(f"Migration QA failed: {item['domain']} {item['viewport']}")
        layout = item.get("layout", {})
        if layout.get("documentWidth", 0) > layout.get("viewportWidth", 0):
            raise ValueError(f"Migration overflow: {item['domain']} {item['viewport']}")
    placeholders = {item["domain"] for item in migration_results if item.get("sourcePlaceholder")}
    if placeholders != {"othodi-spb.ru"}:
        raise ValueError(f"Unexpected source placeholders: {sorted(placeholders)}")

    mchs_results = load_json(MCHS_VRN_RESULTS)
    if len(mchs_results) != 2:
        raise ValueError("Expected desktop/mobile mchs-vrn staged QA")
    for item in mchs_results:
        layout = item.get("layout", {})
        if item.get("status") != 200 or item.get("pageErrors"):
            raise ValueError(f"mchs-vrn staged QA failed: {item['viewport']}")
        if layout.get("documentWidth") != layout.get("viewportWidth"):
            raise ValueError(f"mchs-vrn staged overflow: {item['viewport']}")

    return {
        "fresh_views": len(result_index),
        "fresh_screenshots": screenshot_count,
        "accepted_submissions": len(delivery_index),
        "recipient_routes": recipient_summary.get("passed"),
        "mailbox_confirmed_messages": len(marker_hits),
        "mailbox_confirmed_sites": list(MAILBOX_CONFIRMED_SITES),
        "hidden_background_videos": hidden_video_summary.get("passed"),
        "migration_views": len(migration_results),
        "mchs_vrn_staged_views": len(mchs_results),
    }


def build_cover_note() -> None:
    doc = Document()
    configure_document(doc, title="ГК «АП-Риал» | Сопроводительное письмо")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run("СОПРОВОДИТЕЛЬНОЕ ПИСЬМО"), size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    set_run_font(p.add_run("К отчёту после повторной проверки"), size=22, bold=True, color=DARK_BLUE)

    for text in COVER_NOTE_PARAGRAPHS:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(10 if text != COVER_NOTE_PARAGRAPHS[-1] else 0)
    COVER_NOTE_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(COVER_NOTE_DOCX)


def build_internal_note() -> None:
    content = """# Внутренний разбор инцидента АП-Риал

Дата: 05.08.2026

## Что произошло

Предыдущая система преждевременно закрывала многосайтовые задачи по частичным техническим сигналам. Доступность страницы или успешный отдельный сценарий ошибочно использовались как подтверждение всего поручения. В результате клиент получил отчёт о завершении до полной проверки всех сайтов и всех пунктов.

## Корневые причины

1. Шлюз готовности проверял HTTP и отдельные действия, но не обязательную матрицу «требование × сайт».
2. Визуальные снимки могли быть созданы без обязательного человеческого просмотра и сравнения с формулировкой клиента.
3. Частично выполненное многосайтовое поручение могло закрываться как целое.
4. Почтовая доставка и внешний вид не всегда перепроверялись после последней общей публикации.
5. Автоматизация допускала внешнее сообщение клиенту без отдельного разрешения владельца.
6. Служебные и внешние блокеры смешивались с невыполненными внутренними работами.

## Что изменено

1. Введён обязательный шлюз владельца до начала реализации и отдельный шлюз перед любым сообщением клиенту.
2. Каждое поручение раскладывается на атомарные требования и полный список сайтов.
3. Каждая пара «требование × сайт» обязана иметь статус passed либо явное approved exclusion и собственное доказательство.
4. Каждая фронтенд-правка проверяется зрением на свежих desktop/mobile снимках.
5. Фактическая доставка указывается только для ящика, который действительно открыт и проверен. Принятие заявки обработчиком не выдаётся за доставку письма.
6. Отчёт строится из машинного аудита и фактических снимков; открытые позиции показываются отдельно.
7. Автономный цикл не создаёт черновики и не пишет клиенту. Внешняя коммуникация возможна только по отдельной команде владельца.

## Критерий закрытия

Задача закрывается только после полной матрицы, свежего функционального и визуального QA, проверки доставки, отчёта с доказательствами и решения владельца о выпуске.
"""
    INTERNAL_NOTE.parent.mkdir(parents=True, exist_ok=True)
    INTERNAL_NOTE.write_text(content, encoding="utf-8")


def build_report() -> dict[str, Any]:
    result_index = load_result_index()
    delivery_index = load_delivery_index()
    checks = validate_inputs(result_index, delivery_index)

    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)

    assets = build_human_report_assets(result_index)

    doc = Document()
    configure_document(doc, title="ГК «АП-Риал» | Итоговый отчёт по сайтам")
    properties = doc.core_properties
    properties.title = "Итоговый отчёт о работах по сайтам ГК «АП-Риал»"
    properties.subject = "Исправления после замечаний по формам и сайтам"
    properties.author = "Никита Тихомиров"
    properties.last_modified_by = "Никита Тихомиров"
    properties.keywords = "АП-Риал, сайты, перенос, формы, проверка, отчёт"
    properties.comments = ""

    add_human_report_cover(doc)
    add_human_corrections(doc)
    add_human_visual_evidence(doc, assets)
    add_human_mail_and_video(doc, assets)
    add_human_domain_list(doc)
    add_human_open_items(doc)
    doc.save(OUTPUT_DOCX)

    build_cover_note()
    build_internal_note()

    result = {
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": "docx_generated",
        "client_contact_performed": False,
        "artifacts": {
            "client_report_docx": str(OUTPUT_DOCX.relative_to(ROOT)),
            "client_report_pdf": str(OUTPUT_PDF.relative_to(ROOT)),
            "cover_note_docx": str(COVER_NOTE_DOCX.relative_to(ROOT)),
            "cover_note_pdf": str(COVER_NOTE_PDF.relative_to(ROOT)),
            "internal_note": str(INTERNAL_NOTE.relative_to(ROOT)),
        },
        "checks": checks,
        "forms": {
            "included_domains": len(evidence.INCLUDED_DOMAINS),
            "excluded_domains": list(evidence.EXCLUDED_DOMAINS),
            "requirements": len(FORM_REQUIREMENTS),
            "visual_boards": 0,
            "client_evidence_images": len(assets),
        },
        "migration": {
            "direct_scope": 35,
            "live": MIGRATIONS_LIVE,
            "source_placeholder": MIGRATIONS_SOURCE_PLACEHOLDER,
            "staged_waiting_domain": MIGRATIONS_STAGED,
            "blocked": MIGRATIONS_BLOCKED,
            "scope_decision": MIGRATIONS_SCOPE_DECISION,
            "excluded": MIGRATIONS_EXCLUDED,
            "additional": MIGRATIONS_ADDITIONAL,
            "visual_sheets": 0,
        },
        "client_inputs_required": [item[0] for item in CLIENT_INPUT_REQUIRED],
        "docx_sha256": file_sha256(OUTPUT_DOCX),
        "cover_note_docx_sha256": file_sha256(COVER_NOTE_DOCX),
    }
    OUTPUT_AUDIT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    return result


def finalize_audit(visual_review_manifest: Path | None = None) -> dict[str, Any]:
    from pypdf import PdfReader

    audit = load_json(OUTPUT_AUDIT)
    required = [OUTPUT_DOCX, OUTPUT_PDF, COVER_NOTE_DOCX, COVER_NOTE_PDF, INTERNAL_NOTE]
    for path in required:
        if not path.exists():
            raise FileNotFoundError(path)
    client_report_pages = len(PdfReader(str(OUTPUT_PDF)).pages)
    cover_note_pages = len(PdfReader(str(COVER_NOTE_PDF)).pages)
    visual_review: dict[str, Any] = {
        "client_report_pages": client_report_pages,
        "cover_note_pages": cover_note_pages,
        "all_pages_reviewed": False,
        "result": "pending",
        "reviewed_at": None,
        "reviewer": None,
        "manifest": None,
    }
    if visual_review_manifest is not None:
        review = load_json(visual_review_manifest)
        expected_client_pages = set(range(1, client_report_pages + 1))
        expected_cover_pages = set(range(1, cover_note_pages + 1))
        reviewed_client_pages = set(review.get("client_report_pages", []))
        reviewed_cover_pages = set(review.get("cover_note_pages", []))
        if reviewed_client_pages != expected_client_pages:
            raise ValueError("Visual review manifest does not cover every client report page")
        if reviewed_cover_pages != expected_cover_pages:
            raise ValueError("Visual review manifest does not cover every cover-note page")
        if not review.get("reviewer") or not review.get("reviewed_at"):
            raise ValueError("Visual review manifest must record reviewer and reviewed_at")
        try:
            manifest_label = str(visual_review_manifest.relative_to(ROOT))
        except ValueError:
            manifest_label = str(visual_review_manifest)
        visual_review.update(
            {
                "all_pages_reviewed": True,
                "result": "passed",
                "reviewed_at": review["reviewed_at"],
                "reviewer": review["reviewer"],
                "manifest": manifest_label,
            }
        )

    audit["status"] = "verified_artifacts" if visual_review["all_pages_reviewed"] else "pending_visual_review"
    audit["finalized_at"] = datetime.now().astimezone().isoformat(timespec="seconds")
    audit["artifact_hashes"] = {
        str(path.relative_to(ROOT)): file_sha256(path) for path in required
    }
    audit["pdf_renderer"] = "reportlab"
    audit["visual_review"] = visual_review
    audit["recipient_control"] = {
        "personal_gmail_in_production_recipients": False,
        "configured_recipient_routes": "48/48 matched the client matrix",
        "client_mailbox_delivery_evidence": "2 form messages in info@medlic.spb.ru only",
    }
    OUTPUT_AUDIT.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    return audit


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--finalize-audit", action="store_true")
    parser.add_argument("--visual-review-manifest", type=Path)
    args = parser.parse_args()
    result = finalize_audit(args.visual_review_manifest) if args.finalize_audit else build_report()
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
