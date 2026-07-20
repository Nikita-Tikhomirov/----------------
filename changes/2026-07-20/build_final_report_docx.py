#!/usr/bin/env python3
"""Build the client-facing DOCX report from the verified Markdown source."""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "Финальный-отчет-проверки-2026-07-20.md"
OUTPUT = ROOT / "Отчет-по-формам-и-заявкам-2026-07-20.docx"

INK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 108)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "E8EEF5"
TABLE_WIDTHS_DXA = (1650, 1950, 4000, 1760)


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    properties = cell._tc.get_or_add_tcPr()
    node = properties.find(qn("w:tcW"))
    if node is None:
        node = OxmlElement("w:tcW")
        properties.append(node)
    node.set(qn("w:w"), str(width))
    node.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (
        ("top", top),
        ("bottom", bottom),
        ("start", start),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in TABLE_WIDTHS_DXA:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for cell, value in zip(row.cells, TABLE_WIDTHS_DXA):
            set_cell_width(cell, value)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    return paragraph, run


def add_label_paragraph(doc, label, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.10
    label_run = paragraph.add_run(f"{label}: ")
    set_font(label_run, bold=True)
    value_run = paragraph.add_run(text)
    set_font(value_run)
    return paragraph


def parse_site_rows(source):
    rows = []
    in_table = False
    for line in source.splitlines():
        if line.startswith("| Сайт | Получатель |"):
            in_table = True
            continue
        if in_table and line.startswith("| ---"):
            continue
        if in_table and not line.startswith("|"):
            break
        if in_table:
            parts = [part.strip() for part in line.strip().strip("|").split("|")]
            if len(parts) == 4:
                rows.append([re.sub(r"`([^`]+)`", r"\1", part) for part in parts])
    return rows


def add_site_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ("Сайт", "Получатель", "Выполнено", "Результат")
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_font(run, size=9, bold=True, color=INK)
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(repeat)

    for values in rows:
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, values)):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            if index in (0, 1, 3):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            set_font(run, size=8.5, bold=index == 0)
    set_table_geometry(table)


def add_callout(doc, label, value):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_width(cell, 9360)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    paragraph = cell.paragraphs[0]
    label_run = paragraph.add_run(f"{label}: ")
    set_font(label_run, size=11, bold=True, color=INK)
    value_run = paragraph.add_run(value)
    set_font(value_run, size=11, bold=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def build():
    source = SOURCE.read_text(encoding="utf-8")
    rows = parse_site_rows(source)
    if len(rows) != 30:
        raise RuntimeError(f"Expected 30 site rows, found {len(rows)}")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = INK
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("АП-Риал | Отчет по формам и заявкам")
    set_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Страница ")
    set_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("ОТЧЕТ ПО ФОРМАМ И ЗАЯВКАМ")
    set_font(run, size=22, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run(
        "Завершение унификации и проверки сайтов после переноса на Beget"
    )
    set_font(run, size=12, color=MUTED)

    add_label_paragraph(doc, "Дата", "20 июля 2026 года")
    add_label_paragraph(doc, "Объем", "30 применимых сайтов; 5 исключений клиента")
    add_label_paragraph(
        doc,
        "Статус",
        "Формы опубликованы, визуально проверены и протестированы отправкой",
    )
    add_callout(doc, "Итого к оплате", "74 500 ₽")

    add_heading(doc, "Что приведено к единому стандарту", 1)
    doc.add_paragraph(
        "На применимых сайтах доступны две формы: «ЗАКАЗАТЬ ЗВОНОК» "
        "с телефоном и простой капчей, а также «ЗАДАТЬ ВОПРОС» "
        "с e-mail, необязательным вопросом и простой капчей."
    )
    doc.add_paragraph(
        "Добавлены согласие на обработку персональных данных, ссылка на "
        "политику, заметное закрытие, единый текст успешной отправки, "
        "серверная проверка и честный вывод ошибок."
    )
    doc.add_paragraph(
        "Получатели не заменялись на Gmail: заявки направляются в штатный "
        "ящик каждого сайта."
    )

    add_heading(doc, "Результат по каждому сайту", 1)
    add_site_table(doc, rows)

    add_heading(doc, "Исключения клиента", 1)
    doc.add_paragraph(
        "Формы не добавлялись и не изменялись на rectavr.ru, fstek.spb.ru, "
        "lic-k.ru, apreal-samara.ru и ed-krd.ru."
    )

    add_heading(doc, "Проверки", 1)
    for label, text in (
        ("Публичная доступность", "30 из 30 сайтов отвечают HTTP 200."),
        (
            "Визуальная проверка",
            "30 из 30 сайтов прошли отдельный Chromium-прогон: обе формы, "
            "закрытие, поля, политика и мобильный экран 390×844.",
        ),
        (
            "Реальные заявки",
            "Через каждый из 22 новых обработчиков отправлена маркированная "
            "заявка; 22 из 22 приняты почтовым транспортом.",
        ),
        (
            "Форма вопроса",
            "Валидация проверена на всех 22 обработчиках; реальные вопросы "
            "дополнительно приняты docp.ru и fste.ru.",
        ),
        (
            "Сервер",
            "18 WordPress-компонентов прошли php -l; контрольные суммы "
            "серверных и локальных файлов совпали.",
        ),
    ):
        add_label_paragraph(doc, label, text)

    add_heading(doc, "Дополнительные исправления", 1)
    doc.add_paragraph(
        "Устранены проблемы старой кодировки на fste.ru и lfsb.ru, "
        "наложение скрытых модальных окон, перекрытие кнопок чатами и "
        "сдвиг форм старой темой nousro-nn.ru."
    )
    doc.add_paragraph(
        "На shopap.ru безопасное подключение выполнено через шаблоны "
        "футера OpenCart; сайт повторно проверен после восстановления "
        "исходного index.php."
    )
    doc.add_paragraph(
        "Спам apreal-volgograd.ru шел через старую CF7-форму №3317 без "
        "капчи. Три устаревшие CF7-формы отключены. Прямой бот-POST теперь "
        "возвращает status: spam, новая форма продолжает принимать заявки."
    )

    add_heading(doc, "Почта и DKIM", 1)
    doc.add_paragraph(
        "MX и SPF у apreal36.ru и mchs-spb.ru настроены. Присланный клиентом "
        "DKIM относится к dpomuc.ru и не может использоваться для других "
        "доменов. Нужны два отдельных значения из VK Workspace после выбора "
        "сначала apreal36.ru, затем mchs-spb.ru."
    )

    add_heading(doc, "Оплата", 1)
    add_label_paragraph(doc, "Перенос сайтов", "45 000 ₽")
    add_label_paragraph(doc, "Ранее согласованные исправления", "29 500 ₽")
    add_callout(doc, "Общая сумма", "74 500 ₽")
    doc.add_paragraph(
        "Отдельная сумма за завершающую унификацию форм не начислялась. "
        "Клиент сообщил, что бухгалтерия пришлет реквизиты для счета "
        "21 июля 2026 года."
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
